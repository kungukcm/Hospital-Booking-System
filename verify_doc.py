import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document("AI_Patient_Support_Thesis_Kung'u_Kelvin_Mathigi_2026.docx")
paras = [p.text for p in doc.paragraphs if p.text.strip()]
print(f"Total paragraphs: {len(paras)}")

size_kb = os.path.getsize("AI_Patient_Support_Thesis_Kung'u_Kelvin_Mathigi_2026.docx") // 1024
print(f"File size: {size_kb} KB")

total_words = sum(len(p.split()) for p in paras)
print(f"Estimated word count: {total_words:,}")

print("\nChapter headings:")
for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        print("  [CH]", p.text[:80])
