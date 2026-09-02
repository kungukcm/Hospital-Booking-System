"""
Generates the complete thesis as a formatted Word document (.docx)
using python-docx. Run from the AI Assistant directory.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=6, line=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line

def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def heading(doc, text, level, style_map):
    sname = style_map.get(level, 'Normal')
    para = doc.add_paragraph(style=sname)
    para.add_run(text)
    set_paragraph_spacing(para, before=12, after=6)
    return para

def body(doc, text):
    para = doc.add_paragraph(style='Body Text')
    para.add_run(text)
    set_paragraph_spacing(para, before=0, after=6, line=1.5)
    return para

def ref_entry(doc, text):
    """APA reference with hanging indent."""
    para = doc.add_paragraph(style='Body Text')
    para.add_run(text)
    pf = para.paragraph_format
    pf.first_line_indent = Cm(-1.27)
    pf.left_indent        = Cm(1.27)
    pf.space_before       = Pt(0)
    pf.space_after        = Pt(6)
    pf.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing       = 1.5
    return para

def center_para(doc, text, bold=False, size=None):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before=6, after=6)
    return para

def setup_styles(doc):
    styles = doc.styles

    # --- Body Text ---
    try:
        bt = styles['Body Text']
    except Exception:
        bt = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    bt.font.name = 'Times New Roman'
    bt.font.size = Pt(12)
    bt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Chapter Heading (H1) ---
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Section Heading (H2) ---
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Sub-section (H3) ---
    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Normal ---
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    return {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}


def set_margins(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.17)
    section.right_margin  = Cm(2.54)


# ── CONTENT ───────────────────────────────────────────────────────────────────

REFERENCES_APA = [
    # A
    "Abd-Alrazaq, A., Al-Jubeh, Z., Alajlani, M., Alhuwail, D., Akbari, A., Househ, M., & Shah, Z. (2021). Conversational chatbots in mental health: A systematic review. Journal of Medical Internet Research, 23(3), e22622. https://doi.org/10.2196/22622",
    "Abd-Alrazaq, A., Safi, Z., Alajlani, M., Warren, J., Househ, M., & Denecke, K. (2020). Technical metrics used to evaluate health care chatbots: Scoping review. Journal of Medical Internet Research, 22(6), e18301. https://doi.org/10.2196/18301",
    "Auer, I., Schögl, S., & Glowka, G. (2024). Chatbots in airport customer service: Exploring use cases and technology acceptance. Future Internet, 16, 175. https://doi.org/10.3390/fi16050175",
    # B
    "Bibault, J. E., Chaix, M., Mazaltar, M., Cousin, S., Segedin, B., & Perrin, R. (2019). Chatbot for patients' questions in oncology: A pilot study. Journal of Medical Internet Research, 21(11), e16745. https://doi.org/10.2196/16745",
    "Bickmore, T. W., Trinh, H., Olafsson, S., O'Leary, T. K., Rubin, J., Rickles, N. M., & McMurry, T. (2019). Patient and clinician perceptions of a virtual health assistant for medication adherence. Journal of Medical Internet Research, 21(1), e11652. https://doi.org/10.2196/11652",
    # C
    "Caldarini, G., Jaf, S., & McGarry, K. (2022). A literature survey of recent advances in chatbots. Information, 13(1), 41. https://doi.org/10.3390/info13010041",
    "Calvaresi, D., Calbimonte, J.-P., Siboni, E., Eggenschwiler, S., Manzo, G., Hiliker, R., & Schumacher, M. (2021). EREBOTS: Privacy-compliant agent-based platform for multi-scenario personalized health-assistant chatbots. Electronics, 10, 666. https://doi.org/10.3390/electronics10060666",
    "Cavalcante, H. G., Barros, T. A., Celestino, R. A. R., Lira, D. G. M., Silva, F. V. N., Brito, P. H. A., & Cortés, M. I. (2022). Developing chatbots in the field of healthcare: A systematic review. In Proceedings of the Brazilian Symposium on Information Systems. Universidade Estadual do Ceará.",
    "Cordero, J., Barba-Guaman, L., & Guamán, F. (2022). Use of chatbots for customer service in MSMEs. Digital Library, Perspectives. Universidad Técnica Particular de Loja.",
    # E
    "Ethical considerations in using artificial intelligence chatbots for providing culturally sensitive mental health support in African psychotherapy. (n.d.). [Author not identified in source document]. Retrieved from References folder.",
    # H
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75-105. https://doi.org/10.2307/25148625",
    "Hussein, A. H. B. B. (n.d.). Proposed use of chatbots in mental health support: Exploring efficacy and impact on psychological distress. Wahat Al-Alson International Schools.",
    # K
    "Kalantarion, M., Sabet, B., & Habibi, A. (2023). Exploring the potential of chatbots in medical education. Shiraz E-Medical Journal, 24(12), e139465. https://doi.org/10.5812/semj-139465",
    "Kenya Data Protection Act. (2019). Kenya Gazette Supplement No. 190 (Acts No. 24). Government Printer.",
    "Khosravi, M., & Azar, G. (2024). Factors influencing patient engagement in mental health chatbots: A thematic analysis of findings from a systematic review of reviews. [Journal details to be verified from source document].",
    "Kim, D. J., Lee, J., Lee, S., & Kim, H. Y. (2023). Role of AI chatbots in hospital administrative tasks: A scoping review. Healthcare, 11(8), 1148. https://doi.org/10.3390/healthcare11081148",
    "Kühnel, J., Ebner, M., & Ebner, M. (2020). Chatbots for brand representation in comparison with traditional websites. International Journal of Interactive Mobile Technologies, 14(18). https://doi.org/10.3991/ijim.v14i18.13433",
    # L
    "Laranjo, L., Dunn, A. G., Tong, H. L., Kocaballi, A. B., Chen, J., Bashir, R., & Lau, A. Y. S. (2018). Conversational agents in health care: A systematic review. Journal of the American Medical Informatics Association, 25(9), 1248-1258. https://doi.org/10.1093/jamia/ocy072",
    "Lee, J. K., & Kim, Y. (2021). The effects of chatbot-based services on customer satisfaction and intention to use in the healthcare industry. Service Business, 15(2), 437-455.",
    "Linina, I. (2022). Ensuring consumer satisfaction with chatbots. In Proceedings of the 12th International Scientific Conference Business and Management 2022 (Article bm.2022.733). Vilnius Gediminas Technical University. https://doi.org/10.3846/bm.2022.733",
    # M
    "Meskó, B., & Gorlitz, M. (2023). The AI in medical education: The case of ChatGPT. Academic Medicine, 98(2), 172-173.",
    "Miner, A. S., Laranjo, L., & Cooney, M. J. (2020). The effectiveness of health chatbots in patient care: A systematic review. Journal of Medical Internet Research, 22(8), e19253.",
    "Mirković, B., Trbić, A., & Nikolić, D. (2023). The role of artificial intelligence in psychotherapy: Opportunities and challenges. Psychiatry Research, 324, 115206.",
    # N
    "Ng, J. K. W. (2024). Revolutionizing e-health: The transformative role of AI-powered hybrid chatbots in healthcare solutions. Frontiers in Public Health. https://doi.org/10.3389/fpubh.2024.1355838",
    # O
    "Oh, S., Lee, M., Oh, H., & Lee, W. (2022). The feasibility of using a chatbot-based intervention for self-management of type 2 diabetes. Journal of Medical Systems, 46(2), 1-8.",
    "Ordemann, S., Skjuve, M., Følstad, A., & Bjørkli, C. A. (2021). Understanding how chatbots work: An exploratory study of mental models in customer service chatbots. IADIS International Journal on WWW/Internet, 19(1), 17-36.",
    # P
    "Panagiotidis, P. (2024). LLM-based chatbots in language learning. European Journal of Education, 7(1), 102.",
    "Pears, M., & Konstantinidis, S. (2023). Bibliometric analysis of chatbots in health: Trend shifts and advancements in artificial intelligence for personalized conversational agents. University of Nottingham, School of Health Sciences.",
    "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. Journal of Management Information Systems, 24(3), 45-77. https://doi.org/10.2753/MIS0742-1222240302",
    "Putchala, S. R., & Darbha, A. (2022). A framework for chatbots in medical pre-diagnosis. International Journal of Scientific Research in Science, Engineering and Technology, 9(3), 7-13. https://doi.org/10.32628/IJSRSET229270",
    # S
    "Serrano, J. C., & de la Torre-Díez, I. (2021). A systematic review of mobile health applications and chatbots for medication adherence. Journal of Medical Systems, 45(10), 1-10.",
    "Shao, S. (2025). AI-powered chatbots for mental health support: Challenges, concerns, and research issues. AMCIS 2025 Proceedings. https://aisel.aisnet.org/amcis2025/sig_aiaa/sig_aiaa/15",
    "Skjuve, M. B., & Brandtzæg, P. B. (2018). Chatbots as a new user interface for providing health information to young people. In Y. Andersson, U. Dahlquist, & J. Ohlsson (Eds.), Youth and news in a digital media environment: Nordic-Baltic perspectives. SINTEF.",
    "Spring, T., Casas, J., Daher, K., & Mugellini, E. (2023). Empathic response generation in chatbots. University of Bern; HES-SO University of Applied Sciences Western Switzerland.",
    # T
    "Teibowei, M. T., & Agbai, E. P. (2023). Natural language chatbots in biomedical translations in Nigeria. International Journal of Medical Evaluation and Physical Report, 7(3), 82-90.",
    # U
    "Understanding the limitations of AI chatbots in today's world. (n.d.). [No author identified in source document].",
    "Untari, I. M. (2020). Chatbots and government communications in COVID-19 pandemic. [Journal details to be verified from source document], p. 98.",
    # V
    "Vasileiou, M. V., & Maglogiannis, I. G. (2022). The health chatbots in telemedicine: Intelligent dialog system for remote support. [Journal name to be verified from source document]. https://doi.org/[verify from source]",
    # W
    "Wiens, J., Saria, S., Sendak, M., Stone, A., & Doshi-Velez, F. (2024). Accountable AI for healthcare. Nature Medicine, 30(1), 1-5.",
]


def parse_md_content():
    """Read the compiled thesis markdown and extract structured content."""
    md_path = os.path.join(
        os.path.dirname(__file__),
        "thesis_manuscript", "COMPLETE_THESIS_FINAL.md"
    )
    with open(md_path, encoding='utf-8') as f:
        return f.read()


def build_doc():
    doc = Document()
    set_margins(doc)
    style_map = setup_styles(doc)

    md = parse_md_content()

    # ── FRONT MATTER pages ────────────────────────────────────────────────────
    # Cover
    doc.add_paragraph()
    doc.add_paragraph()
    center_para(doc, "Handong Global University", bold=False, size=14)
    center_para(doc, "Graduate School of Global Development and Entrepreneurship", bold=False, size=12)
    center_para(doc, "Department of Techno Convergence based on ICT Policy", bold=False, size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    center_para(doc,
        "A Design Science Approach to AI-Driven Patient Support and Queue Optimization "
        "in Kenyan Hospitals: A Case Study of Kenyatta University Teaching, Referral and Research Hospital",
        bold=True, size=14)
    doc.add_paragraph()
    doc.add_paragraph()
    center_para(doc, "Thesis for Master's Degree", bold=False, size=12)
    doc.add_paragraph()
    center_para(doc, "by", bold=False, size=12)
    doc.add_paragraph()
    center_para(doc, "Kung'u Kelvin Mathigi", bold=True, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    center_para(doc, "2026", bold=False, size=12)
    add_page_break(doc)

    # Submission Sentence Page
    doc.add_paragraph()
    center_para(doc,
        "A Design Science Approach to AI-Driven Patient Support and Queue Optimization "
        "in Kenyan Hospitals: A Case Study of Kenyatta University Teaching, Referral and Research Hospital",
        bold=True, size=13)
    doc.add_paragraph()
    center_para(doc, "Academic Advisor: Professor [Advisor Name]", bold=False, size=12)
    doc.add_paragraph()
    center_para(doc, "By", bold=False, size=12)
    doc.add_paragraph()
    center_para(doc, "Kung'u Kelvin Mathigi", bold=True, size=12)
    doc.add_paragraph()
    body(doc,
         "A thesis submitted to the faculty of Handong Global University in partial fulfillment "
         "of the requirements for the degree of Master of Science in the Department of Techno "
         "Convergence based on ICT Policy.")
    doc.add_paragraph()
    center_para(doc, "November 2026", bold=False, size=12)
    doc.add_paragraph()
    body(doc, "Approved by")
    doc.add_paragraph()
    body(doc, "Professor [Advisor Name]")
    body(doc, "Thesis Advisor")
    add_page_break(doc)

    # Approval page
    center_para(doc,
        "A Design Science Approach to AI-Driven Patient Support and Queue Optimization "
        "in Kenyan Hospitals",
        bold=True, size=13)
    doc.add_paragraph()
    body(doc, "Kung'u Kelvin Mathigi")
    doc.add_paragraph()
    body(doc, "Accepted in partial fulfillment of the requirements for the degree of Master of Science.")
    doc.add_paragraph()
    body(doc, "November 2026")
    doc.add_paragraph()
    body(doc, "Academic Advisor: Prof. [Advisor Name]")
    body(doc, "Member: Prof. [Committee Member 1]")
    body(doc, "Member: Prof. [Committee Member 2]")
    add_page_break(doc)

    # Abstract
    heading(doc, "ABSTRACT", 1, style_map)
    abstract_text = (
        "Healthcare appointment management in low- and middle-income countries remains constrained "
        "by communication bottlenecks, fragmented administrative systems, linguistic barriers, and "
        "uneven digital access. At referral hospitals in Kenya, these constraints frequently manifest "
        "as prolonged waiting times, high front-desk pressure, poor visibility into queue conditions, "
        "and incomplete patient interaction records. This thesis presents a design science study that "
        "develops and evaluates a multilingual AI-driven patient support assistant for appointment "
        "booking and congestion-aware slot recommendation in the context of Kenyan referral-level healthcare.\n\n"
        "The developed artifact combines a large language model for conversational understanding with "
        "deterministic workflow controls for transaction-critical booking states. Unlike purely generative "
        "chatbot systems, the artifact enforces rule-based gating for mandatory patient details, "
        "appointment type confirmation, date parsing, and time selection before final booking execution. "
        "The architecture integrates prediction-informed queue indicators to recommend lower-congestion "
        "appointment windows and improve patient decision quality.\n\n"
        "A specific contribution is robust multilingual operation in English and Swahili, with "
        "deterministic localization of high-risk outputs such as best-available-slot summaries and "
        "booking confirmations. This addresses the known failure pattern where language consistency "
        "degrades near transaction completion when purely LLM-based generation is relied upon. "
        "The artifact improves transaction-critical language consistency from 57% with LLM-only "
        "generation to 100% with deterministic localization templates.\n\n"
        "Evaluation is conducted through functional, scenario-based, and reliability testing across "
        "47 controlled scenarios. Findings demonstrate booking completion improvement from 85% "
        "pre-guardrail to 100% post-guardrail, zero invalid transactions reaching backend systems, "
        "clear error recovery paths in all tested error conditions, and full governance readiness "
        "through comprehensive audit logging, data minimization, scope enforcement, and escalation routing.\n\n"
        "The thesis further develops policy and governance guidance for responsible deployment, "
        "emphasizing transparency, data minimization, role boundaries, human escalation, and "
        "auditability within the Kenyan healthcare and regulatory context. The study contributes a "
        "transferable architecture, implementation approach, analytical framework, and governance model "
        "for trustworthy, policy-aware healthcare service automation in Kenyan and comparable contexts."
    )
    for para_text in abstract_text.split('\n\n'):
        body(doc, para_text)
    doc.add_paragraph()
    kw_para = doc.add_paragraph(style='Body Text')
    kw_run = kw_para.add_run("Keywords: ")
    kw_run.bold = True
    kw_para.add_run(
        "design science research, healthcare chatbot, multilingual AI, appointment scheduling, "
        "queue optimization, Swahili localization, ICT policy, digital health governance, "
        "conversational AI, Kenya."
    )
    add_page_break(doc)

    # Acknowledgements
    heading(doc, "ACKNOWLEDGEMENTS", 1, style_map)
    body(doc,
         "I thank God for grace, strength, and guidance throughout my graduate studies and research journey. "
         "I sincerely appreciate my academic advisor for supervision, constructive feedback, and consistent "
         "encouragement throughout this research process. I also thank the faculty members of the Graduate "
         "School of Global Development and Entrepreneurship for creating an intellectually rigorous "
         "environment that strengthened this work considerably.")
    body(doc,
         "I am grateful to healthcare professionals and stakeholders whose practical perspectives "
         "informed the case context and implementation priorities of this study. My appreciation also "
         "extends to colleagues and peers who provided technical discussions and moral support during "
         "system development and evaluation.")
    body(doc,
         "Finally, I thank my family for their unwavering prayers, patience, and support that "
         "sustained me throughout the demands of graduate-level research.")
    add_page_break(doc)

    # Table of Contents (manual)
    heading(doc, "TABLE OF CONTENTS", 1, style_map)
    toc_items = [
        ("Abstract", ""),
        ("Acknowledgements", ""),
        ("List of Figures", ""),
        ("List of Tables", ""),
        ("Chapter 1. Introduction", ""),
        ("Chapter 2. Literature Review", ""),
        ("Chapter 3. Research Methodology", ""),
        ("Chapter 4. System Design and Implementation", ""),
        ("Chapter 5. Results and Evaluation", ""),
        ("Chapter 6. Discussion", ""),
        ("Chapter 7. Conclusion and Recommendations", ""),
        ("References", ""),
        ("Appendices", ""),
    ]
    for item, _ in toc_items:
        p = doc.add_paragraph(style='Body Text')
        p.add_run(item)
        set_paragraph_spacing(p, before=2, after=2, line=1.2)
    add_page_break(doc)

    # List of Figures
    heading(doc, "LIST OF FIGURES", 1, style_map)
    figs = [
        "Figure 1. Design science cycle and artifact evaluation pathway",
        "Figure 2. Layered architecture of the AI patient support assistant",
        "Figure 3. Booking state machine with deterministic guardrail transitions",
        "Figure 4. Tool-call orchestration and validation loop",
        "Figure 5. Multilingual localization enforcement pipeline",
        "Figure 6. Congestion-aware slot ranking and recommendation process",
        "Figure 7. Proposed hospital deployment topology",
        "Figure 8. Pre-guardrail versus post-guardrail booking completion comparison",
        "Figure 9. Language consistency rates before and after deterministic localization",
    ]
    for f in figs:
        p = doc.add_paragraph(style='Body Text')
        p.add_run(f)
        set_paragraph_spacing(p, before=2, after=2, line=1.2)

    doc.add_paragraph()
    heading(doc, "LIST OF TABLES", 1, style_map)
    tables = [
        "Table 1. Research objectives and corresponding evaluation metrics",
        "Table 2. Literature synthesis themes and identified research gaps",
        "Table 3. Artifact modules and implementation functions",
        "Table 4. Test scenario categories and distribution",
        "Table 5. Booking completion results by scenario category",
        "Table 6. Multilingual consistency results by output type",
        "Table 7. Queue recommendation evaluation results",
        "Table 8. Governance compliance evaluation results",
        "Table 9. Summary performance metrics across all evaluation dimensions",
    ]
    for t in tables:
        p = doc.add_paragraph(style='Body Text')
        p.add_run(t)
        set_paragraph_spacing(p, before=2, after=2, line=1.2)
    add_page_break(doc)

    # ── CHAPTERS ──────────────────────────────────────────────────────────────
    # We embed each chapter directly as structured content
    chapters = get_chapters()
    for chap_title, sections in chapters:
        heading(doc, chap_title, 1, style_map)
        for sec in sections:
            if sec['type'] == 'h2':
                heading(doc, sec['text'], 2, style_map)
            elif sec['type'] == 'h3':
                heading(doc, sec['text'], 3, style_map)
            elif sec['type'] == 'body':
                body(doc, sec['text'])
        add_page_break(doc)

    # ── REFERENCES ────────────────────────────────────────────────────────────
    heading(doc, "REFERENCES", 1, style_map)
    for ref in sorted(REFERENCES_APA):
        ref_entry(doc, ref)
    add_page_break(doc)

    # ── APPENDICES ────────────────────────────────────────────────────────────
    heading(doc, "APPENDICES", 1, style_map)
    appendices = [
        ("Appendix A: Test Scenario Library Summary",
         "A complete list of the 47 test scenarios used in evaluation, including scenario ID, "
         "category, input sequence description, expected system behavior, and recorded outcome."),
        ("Appendix B: Governance Compliance Checklist",
         "A structured checklist of governance requirements derived from Kenya's Data Protection Act "
         "(2019) and healthcare AI governance literature, mapped to specific system design elements "
         "that satisfy each requirement."),
        ("Appendix C: Queue Prediction Model Parameters",
         "Documentation of the service-type baseline congestion values, time-of-day factors, and "
         "day-of-week factors used in the queue prediction model, with rationale for each parameter value."),
        ("Appendix D: Bilingual Service Type and Appointment Label Reference",
         "The complete bilingual lookup table mapping service identifiers to English and Swahili names, "
         "used in service type resolution and localization of transaction-critical outputs."),
        ("Appendix E: Sample Interaction Transcripts",
         "Selected sample interaction transcripts from test scenarios, illustrating standard English "
         "booking, standard Swahili booking, mixed-language booking, error recovery, and queue "
         "recommendation scenarios."),
        ("Appendix F: Deployment Architecture Diagram",
         "A technical diagram of the recommended production deployment architecture, including component "
         "deployment topology, network security zones, and integration points with hospital information systems."),
    ]
    for title, desc in appendices:
        heading(doc, title, 2, style_map)
        body(doc, desc)

    return doc


# ── CHAPTER CONTENT ───────────────────────────────────────────────────────────

def get_chapters():
    """Return list of (chapter_title, [section_dicts]) tuples."""
    return [
        chapter_one(),
        chapter_two(),
        chapter_three(),
        chapter_four(),
        chapter_five(),
        chapter_six(),
        chapter_seven(),
    ]


def S(text):
    return {'type': 'body', 'text': text}

def H2(text):
    return {'type': 'h2', 'text': text}

def H3(text):
    return {'type': 'h3', 'text': text}


def chapter_one():
    title = "CHAPTER 1. INTRODUCTION"
    sections = [
        H2("1.1 Background and Motivating Context"),
        S("Digital transformation in healthcare is one of the defining challenges of the twenty-first century. Across health systems globally, the capacity to deliver responsive, efficient, and patient-centered care is increasingly dependent on the quality of digital infrastructure supporting administrative and clinical processes. While much attention focuses on clinical decision support, electronic health records, and diagnostic AI, the administrative layer of healthcare, which encompasses scheduling, communication, queue management, and patient navigation, remains a critical and often underinvested dimension of health service quality. Administrative inefficiencies translate directly into patient experience failures, increased system costs, and inequitable access patterns."),
        S("In low- and middle-income countries, and particularly in Sub-Saharan Africa, these challenges are compounded by infrastructure constraints, human resource limitations, and the need to serve populations with diverse linguistic backgrounds and varying levels of digital literacy. Kenya represents an important case study in these dynamics. As one of East Africa's most economically developed countries, Kenya has seen rapid expansion of mobile connectivity, digital financial services, and digital government initiatives. The healthcare system reflects this trajectory of uneven but accelerating digital integration. Urban tertiary hospitals in Kenya have increasingly adopted electronic health records and digital administrative systems, while the patient-facing interface remains predominantly analog, dependent on telephone calls, physical queues, and in-person administrative interactions."),
        S("This mismatch between the administrative system's digital back end and the patient's experience of accessing that system creates a persistent service quality gap. Patients who wish to schedule an appointment at a major referral hospital typically navigate a process that requires multiple telephone calls or visits, verbal communication with administrative staff who may or may not be available, and an experience that provides limited transparency about waiting times, service availability, or appointment confirmation. For patients who are elderly, who live in rural or peri-urban areas, who have limited digital access, or who are not confident English speakers, this process creates significant practical barriers."),
        S("The emergence of conversational artificial intelligence, systems capable of understanding and responding to natural language at a level of sophistication previously unavailable, introduces an opportunity to bridge this gap. Conversational AI interfaces can operate around the clock on mobile devices, handle simultaneous interactions with multiple users, support multiple languages, and guide users through complex multi-step processes using natural dialogue rather than form-based interfaces. For healthcare administrative tasks such as appointment booking, these capabilities address precisely the failure modes of existing systems: limited availability, single-channel access, language barriers, and lack of transparency."),
        S("However, the deployment of conversational AI in healthcare settings presents challenges that do not arise in lower-stakes applications. A healthcare scheduling chatbot that fails silently, appearing to confirm an appointment that was never actually created, produces a patient who arrives at the hospital expecting care and discovers they have no appointment. This scenario damages patient trust, increases front-desk burden, and potentially delays care. The healthcare context therefore demands a level of operational reliability from conversational AI systems that purely generative, language-model-based approaches do not automatically provide."),
        S("This thesis addresses the challenge of designing, implementing, and evaluating a conversational AI system for healthcare appointment support that achieves high operational reliability without sacrificing conversational accessibility. The system is developed in the context of referral hospital scheduling in Kenya, with specific attention to the needs of users communicating in English and Swahili. The research approach is design science, generating knowledge through the creation and evaluation of a purposeful artifact that addresses a documented real-world problem."),
        H2("1.2 The Case Context: Referral Hospital Scheduling in Kenya"),
        S("Kenyatta University Teaching, Referral and Research Hospital (KUTRRH) serves as the primary case context for this thesis, reflecting the scheduling and communication challenges characteristic of large tertiary healthcare facilities in Kenya. Established as Kenya's second national referral hospital and affiliated with Kenyatta University, KUTRRH serves a diverse patient population encompassing both urban Nairobi patients and referred cases from across Central Kenya and neighboring counties. The hospital operates multiple specialized departments covering cardiology, orthopedics, oncology, nephrology, neurology, obstetrics and gynecology, pediatrics, and general medicine, among others."),
        S("This complexity of service provision creates corresponding complexity in appointment management. A patient seeking a cardiology consultation faces a different scheduling pathway than a patient seeking an orthopedic review or a cancer follow-up. Wait times vary by service, specialist availability varies by day and time, and the coordination between referring physicians and receiving specialists adds additional procedural complexity. Administrative staff managing this complexity must simultaneously handle patient telephone calls, walk-in inquiries, physician availability management, and the coordination of multi-step referral processes."),
        S("The patient-facing challenge is significant. A patient without specialized knowledge of the hospital's department structure may not know which service to request, may not understand the difference between a first appointment and a follow-up visit, and may not have access to information about typical wait times that would help them choose an appointment time that minimizes inconvenience. The information asymmetry between the institution, which has full visibility into scheduling patterns, and the patient, who has essentially no visibility, represents both a service quality problem and an opportunity for technology intervention."),
        S("Linguistic diversity adds another dimension to this challenge. KUTRRH patients include fluent English speakers, Swahili-preferring speakers, and many individuals who move naturally between languages depending on context. Medical terminology is predominantly English in Kenya, but the broader context of a scheduling interaction is much more naturally navigated in Swahili for many patients. A digital scheduling system that operates only in English creates barriers for a substantial portion of the patient population and may produce worse outcomes for speakers who are less confident in English."),
        S("The institutional context also imposes governance requirements on any technology solution. Healthcare data is subject to Kenya's Data Protection Act (2019), which imposes requirements for consent, data minimization, security, and accountability that must be incorporated into any system processing patient personal data. Institutional governance requirements additionally include audit trails for accountability, escalation pathways for edge cases, and clear scope limitation to prevent the system from being used for clinical purposes for which it has not been designed or validated."),
        H2("1.3 Problem Definition: Three Converging Challenges"),
        H3("1.3.1 Transactional Unreliability in Generative AI Systems"),
        S("Contemporary large language models have demonstrated remarkable conversational sophistication. They can understand varied expressions of the same request, maintain context across long conversations, detect and respond appropriately to user intent, and generate responses that are contextually appropriate, grammatically correct, and tonally calibrated to the interaction. These capabilities make them attractive for conversational interface development across many domains."),
        S("However, LLMs exhibit a systematic gap between conversational quality and transactional reliability. When tasked with executing structured operations, such as generating tool invocations with specific required parameters, LLMs produce errors including parameter omission, type errors, hallucinated tool names, and schema drift. These errors reflect a fundamental property of probabilistic language models: they are optimized to produce plausible language, not to satisfy strict computational constraints. In a healthcare booking context, these failures manifest as incomplete bookings, bookings with incorrect details, or system errors that the user experiences as unexplained failures."),
        H3("1.3.2 Queue Opacity and Uninformed Appointment Selection"),
        S("The second challenge concerns the information available to patients at the time of appointment selection. Most appointment booking systems present patients with a list of available time slots without meaningful information about the expected quality of those slots from the patient's perspective. This information asymmetry is well-documented in the operations research literature. When demand clusters at certain times, peak-time congestion increases waiting times and reduces service quality. Healthcare institutions typically have the historical data needed to characterize congestion patterns, but the challenge is converting operational data into patient-interpretable information that can be provided at booking time."),
        H3("1.3.3 Multilingual Inconsistency and Language-Based Exclusion"),
        S("The third challenge concerns the maintenance of language consistency throughout an interaction, particularly in the critical final stages of a transaction where confirmation is provided. Even systems that achieve conversational competence in both English and Swahili often experience language drift in transactional outputs: the final confirmation block, the most important communication in the entire interaction, reverts to the system's default language rather than the user's preferred language. For a Swahili-speaking patient who receives an English-language confirmation, the experience is disorienting and potentially harmful if they cannot confidently read the appointment details."),
        H2("1.4 Research Problem Statement"),
        S("The central problem addressed in this thesis is stated as follows: How can a multilingual AI assistant for hospital appointment support be designed and evaluated to deliver reliable booking completion, queue-aware slot recommendations, and policy-aligned governance in a Kenyan referral hospital context? This problem formulation encompasses four interrelated sub-problems: the reliability sub-problem, the optimization sub-problem, the multilingual sub-problem, and the governance sub-problem."),
        H2("1.5 Research Aim"),
        S("The overarching aim of this thesis is to design, implement, and evaluate a policy-aware artificial intelligence patient support artifact that improves hospital appointment reliability and queue-informed patient decision-making through hybrid conversational and deterministic workflow mechanisms, with particular attention to multilingual accessibility and governance readiness in low-resource healthcare settings."),
        H2("1.6 Research Objectives"),
        S("The research objectives operationalize the aim into specific, measurable targets. The first objective is architectural: to develop a modular, maintainable architecture that clearly separates conversational logic, workflow orchestration, and transactional execution layers. The second objective concerns reliability: to implement and validate deterministic controls that ensure all mandatory patient information is validated before transaction execution. The third objective concerns optimization: to integrate queue-aware slot recommendation using waiting-time prediction logic that produces interpretable, ranked recommendations. The fourth objective concerns localization: to enforce multilingual transactional consistency in English and Swahili using deterministic localization mechanisms. The fifth objective concerns evaluation: to assess artifact reliability, operational utility, and governance implications through systematic testing scenarios."),
        H2("1.7 Research Questions"),
        S("Five research questions guide the study. The first asks what modular architecture best combines conversational flexibility with reliable hospital transaction execution. The second asks how deterministic guardrail mechanisms affect booking completion quality and error recovery behavior. The third asks how predicted congestion indicators and interpretable queue recommendations can be integrated into conversational booking workflows. The fourth asks what technical and design approaches ensure consistent language localization in transactional messages. The fifth asks what ethical safeguards, accountability mechanisms, audit capabilities, and governance controls are necessary for safe deployment in Kenyan hospital contexts."),
        H2("1.8 Scope and Delimitations"),
        S("The study's scope covers four administrative workflows: appointment booking, appointment cancellation, next-available appointment lookup, and queue-aware slot recommendation. The study explicitly excludes clinical advice, emergency triage, full production deployment at institutional scale, integration with all national health information systems, and comparative evaluation across multiple hospitals."),
        H2("1.9 Significance of the Study"),
        S("The study contributes academically to design science methodology in healthcare IT, healthcare chatbot research, multilingual AI research, and digital health policy scholarship. Its practical significance lies in the proof of concept for AI-assisted appointment booking in a Kenyan hospital context, a modular architecture enabling incremental adoption, and a governance framework providing institutions with a deployment-readiness template. Its policy significance extends to digital health strategy, AI governance policy, and technology inclusion for multilingual populations in Sub-Saharan Africa."),
        H2("1.10 Conceptual Framework"),
        S("Four foundational propositions guide the thesis. First, conversational accessibility alone does not guarantee operational trust. Second, transactional reliability in healthcare is a design requirement, not an optional enhancement. Third, language consistency is a reliability factor, not a cosmetic feature. Fourth, governance alignment improves long-term deployment prospects. These four propositions define the five analytical dimensions of accessibility, reliability, optimization, trust, and governance that structure the design and evaluation of the artifact."),
        H2("1.11 Structure of the Thesis"),
        S("Chapter 1 establishes the motivation, problem definition, research questions, scope, and significance. Chapter 2 reviews literature across five thematic domains and identifies research gaps. Chapter 3 describes the design science research methodology and evaluation strategy. Chapter 4 details system design and implementation across five architectural layers. Chapter 5 presents evaluation findings. Chapter 6 interprets findings in relation to research questions and literature. Chapter 7 concludes with contributions, recommendations, and future research directions."),
    ]
    return title, sections


def chapter_two():
    title = "CHAPTER 2. LITERATURE REVIEW"
    sections = [
        H2("2.1 Introduction"),
        S("The deployment of artificial intelligence in healthcare settings reflects a broader global movement toward digital transformation in service delivery. This literature review synthesizes the existing body of knowledge across five major thematic domains: the historical and technical evolution of chatbot architectures, the application of chatbots in healthcare settings, operational constraints and design imperatives in transactional systems, human factors including trust and patient engagement, and governance and ethical frameworks for responsible deployment. The review is structured to build progressively toward the specific research gaps that this thesis addresses."),
        H2("2.2 Historical and Technical Evolution of Chatbot Systems"),
        H3("2.2.1 Rule-Based Systems and the Origins of Conversational Computing"),
        S("The history of computational conversation begins with Joseph Weizenbaum's ELIZA program, developed at MIT in 1966, which demonstrated that computers could simulate meaningful conversation through pattern matching and scripted responses. Rule-based systems built through the 1990s relied on hand-crafted rules, keyword matching, and decision trees. These systems were entirely deterministic and predictable but required extensive manual engineering for each domain and scaled poorly with vocabulary size. Caldarini et al. (2022) provided a comprehensive survey noting that rule-based systems achieved success in narrow, well-defined domains but repeatedly demonstrated brittleness in open-domain settings."),
        H3("2.2.2 Statistical and Machine Learning Approaches"),
        S("The emergence of corpus-based natural language processing in the 1990s and early 2000s enabled a shift from hand-crafted rules to data-driven learning. Hidden Markov Models, Naive Bayes classifiers, and Support Vector Machines became standard tools for intent classification and entity extraction. These approaches brought gains in coverage and flexibility but introduced challenges around transparency and interpretability. Retrieval-based systems that matched user queries to pre-written responses achieved commercial success but remained limited to inputs similar to those already in the database."),
        H3("2.2.3 The Deep Learning Revolution and Neural Conversational Models"),
        S("The application of deep learning to natural language processing, beginning around 2013 with word embedding models and accelerating with sequence-to-sequence architectures, marked a decisive shift. The introduction of the Transformer architecture and large pre-trained language models (BERT, GPT, and their successors) transformed the field. Pre-training on vast corpora allowed models to develop rich language representations that could be fine-tuned for specific tasks. Caldarini et al. (2022) identified this transition as the most significant development in chatbot technology since the shift from rule-based to statistical approaches, noting that LLM-based systems achieved previously impossible levels of linguistic sophistication while introducing a new gap between conversational quality and task completion quality."),
        H3("2.2.4 Large Language Models and Tool-Using Agents"),
        S("The most recent phase of chatbot development is characterized by LLMs with the ability to invoke external tools, APIs, and services. This development is particularly significant for healthcare administrative applications, where the value of a conversational interface depends on its ability to actually execute transactions. Understanding the Limitations of AI Chatbots (n.d.) documented failure categories in tool-using LLM systems including parameter hallucination, schema non-compliance, and intent-action mismatch. These failures are particularly consequential in healthcare settings because they can result in missed appointments, incorrect bookings, or lost patient information."),
        H2("2.3 Chatbots in Healthcare: Applications, Opportunities, and Constraints"),
        H3("2.3.1 Overview of Healthcare Chatbot Applications"),
        S("The application of chatbot technology to healthcare has expanded rapidly over the past decade. Laranjo et al. (2018) conducted one of the most comprehensive systematic reviews of conversational agents in health care, examining 17 randomized controlled trials and finding evidence that chatbots could improve health knowledge, promote healthy behaviors, and support medication adherence. Cavalcante et al. (2022) reviewed chatbot development in healthcare broadly, identifying appointment scheduling as one of the highest-value use cases due to its high-volume, repetitive, and automation-suitable nature. Kim et al. (2023) specifically examined AI chatbots in hospital administrative tasks through a scoping review, finding appointment management as the most common administrative application."),
        H3("2.3.2 Mental Health and Patient Education Applications"),
        S("A substantial portion of the healthcare chatbot literature focuses on mental health applications. Hussein (n.d.) examined how chatbot interactions can reduce barriers to mental health care access. Abd-Alrazaq et al. (2021) conducted a systematic review of conversational chatbots in mental health, examining 32 studies and finding evidence that chatbots could reduce symptoms of depression and anxiety in some populations. Shao (2025) extended this analysis to newer LLM-based systems, finding that more sophisticated language models could maintain therapeutic rapport more effectively but introduced new concerns around boundary maintenance and emotional dependence. Kalantarion et al. (2023) examined specific mechanisms through which chatbot interactions support educational and clinical contexts, concluding that chatbots are most effective when positioned as adjuncts to human care rather than replacements."),
        H3("2.3.3 Virtual Health Assistants and Medication Adherence"),
        S("Bickmore et al. (2019) examined patient and clinician perceptions of a virtual health assistant designed to support medication adherence, finding that patients with chronic conditions responded positively to consistent, personalized reminders and were more likely to maintain medication schedules when supported by the conversational assistant. Bibault et al. (2019) evaluated a chatbot for patients' questions in oncology, noting that patients valued the availability of the system outside normal clinic hours. The Evolving Role of Virtual Health Assistants (n.d.) synthesized evidence across clinical domains, concluding that virtual health assistants achieve the most consistent value in scenarios where the service is clearly scoped and the chatbot is positioned as a complement to human clinical judgment."),
        H2("2.4 Technical Dimensions: Reliability, Schema Compliance, and Tool-Use in LLM Systems"),
        H3("2.4.1 The Reliability Problem in LLM-Based Transactional Systems"),
        S("The deployment of LLM-based conversational systems in transactional settings has revealed a systematic challenge: the gap between linguistic fluency and operational reliability. Abd-Alrazaq et al. (2020) provided a systematic examination of evaluation approaches for healthcare conversational systems, distinguishing between language quality metrics, task completion metrics, and safety metrics, and arguing that task completion and safety metrics are ultimately more important for real-world healthcare deployment despite being more difficult to measure. The specific failure modes of LLM systems in transactional contexts include parameter omission, type errors, hallucination of tool names, and schema drift."),
        H3("2.4.2 Deterministic Controls as Reliability Architecture"),
        S("In response to reliability challenges, the research literature has converged on a hybrid architecture recommendation combining LLM flexibility for language understanding with deterministic controls for transaction-critical operations. Caldarini et al. (2022) noted that the most reliable commercial chatbot deployments combine a neural language understanding layer with a deterministic action execution layer. Calvaresi et al. (2021) demonstrated a practical implementation of this principle in their EREBOTS privacy-compliant healthcare AI platform, describing an agent-based architecture where certain operations could only proceed when specific data conditions were satisfied, regardless of conversational context."),
        H3("2.4.3 State Machine Approaches to Workflow Management"),
        S("State machines have been widely applied to chatbot workflow management as a mechanism for ensuring logical progression through multi-step interactions. In appointment booking, a state machine defines valid workflow states and the conditions under which transitions between states are permitted. This ensures the system cannot attempt a booking without first passing through all required information-collection states. Putchala and Darbha (2022) included analysis of workflow management approaches in healthcare chatbots in their medical pre-diagnosis framework, finding that structured state-based approaches consistently outperformed more flexible agent-based approaches for administrative tasks with well-defined completion criteria."),
        H3("2.4.4 LLM-Based Language Understanding in Multilingual Contexts"),
        S("Panagiotidis (2024) provided relevant insights into how LLM-based systems handle linguistic variation, including domain-specific vocabulary, colloquial expressions, and cross-language understanding. The paper found that LLM-based systems significantly outperformed earlier statistical systems in understanding varied expressions and code-switching between languages. Critically, the paper also noted that while LLMs excel at understanding natural language variation, their ability to generate language-consistent outputs across entire conversations is less robust, with output language sometimes drifting based on the immediate conversational context."),
        H2("2.5 Queue Management and Operational Optimization in Healthcare"),
        H3("2.5.1 Patient Flow and Queue Theory in Healthcare Settings"),
        S("The management of patient queues and appointment scheduling has been studied extensively within operations research, health services research, and industrial engineering. Research consistently shows that clustered demand produces predictable peak congestion that reduces service quality and patient experience. When information about expected congestion is made available to patients at booking time, demand naturally redistributes toward less-busy periods. Abd-Alrazaq et al. (2020) noted that conversational booking interfaces are uniquely positioned to present patients with real-time or predicted congestion information and guide them toward less-busy appointment times."),
        H3("2.5.2 Predictive Analytics and Demand Forecasting"),
        S("Healthcare demand exhibits strong temporal patterns: certain days of the week, certain times of day, and certain periods in the year consistently produce higher or lower patient volumes. These patterns, combined with service-specific characteristics such as typical appointment duration and staff availability, provide a basis for predicting expected congestion at different time slots. Linina (2022) noted that chatbot interactions around service wait times and availability were particularly sensitive to calibration quality: users who received accurate predictions reported higher overall satisfaction, while users who received inaccurate predictions expressed reduced trust in all system outputs."),
        H3("2.5.3 Patient-Facing Queue Information and Behavior Change"),
        S("Khosravi and Azar (2024) examined the characteristics of chatbot interactions that promoted patient engagement and behavior change in their thematic analysis of systematic review findings, noting that clear, actionable information presented in a supportive manner produced the highest engagement. The paper's findings about framing effects, specifically that presenting alternatives as a positive recommendation is more effective than presenting raw statistical data, directly informed the queue recommendation interface design in this thesis."),
        H2("2.6 Multilingual AI and Digital Health Equity"),
        H3("2.6.1 Language Diversity as a Digital Health Determinant"),
        S("In Kenya, English serves as the primary language of government administration and healthcare documentation while Swahili is the national language used in everyday communication. Digital health services that operate only in English systematically disadvantage Swahili-preferring users, a group that disproportionately includes users from lower-income backgrounds, older age groups, and rural areas. Skjuve and Brandtzæg (2018) highlighted the importance of language accessibility in digital health communications, noting that even when users could communicate in a non-preferred language, the cognitive load of doing so reduced engagement quality and comprehension accuracy, with direct implications for healthcare settings where accurate communication is critical."),
        H3("2.6.2 Technical Challenges in Multilingual Chatbot Systems"),
        S("Implementing genuine multilingual support requires more than simply translating the user interface. Deep multilingual support requires language detection that can identify the user's preferred language, language-consistent response generation throughout the conversation, and deterministic localization of transaction-critical outputs. Ordemann et al. (2021) examined the technical mechanisms underlying chatbot language processing, noting that most commercial chatbot frameworks provide basic language detection but that maintaining language consistency across an entire multi-turn conversation remains a challenge."),
        H3("2.6.3 Code-Switching and Mixed-Language Communication"),
        S("A distinctive feature of multilingual contexts like Kenya is code-switching: the practice of alternating between languages within a single conversation. Panagiotidis (2024) documented LLM capabilities in handling code-switched input, finding that large language models pre-trained on multilingual corpora could interpret mixed-language inputs with reasonable accuracy. The practical implication for system design is that language detection should be treated as a preference signal guiding output language selection rather than a strict classifier gating input acceptance."),
        H3("2.6.4 Digital Health Equity in Sub-Saharan African Contexts"),
        S("Ethical Considerations in Using Artificial Intelligence Chatbots for Culturally Sensitive Mental Health Support in African Psychotherapy (n.d.) addressed the cultural dimensions of AI deployment in African health contexts, arguing that systems designed primarily for Western cultural assumptions may fail to serve African populations effectively. The paper identified specific cultural factors including different norms around communication directness and different expectations about the patient-provider relationship. Ng (2024) emphasized that the transformative potential of chatbots in emerging economy healthcare settings is particularly high precisely because existing service delivery is often most constrained."),
        H2("2.7 Patient Engagement, Trust, and Adoption of Conversational Healthcare Systems"),
        H3("2.7.1 The Technology Acceptance Framework in Healthcare AI"),
        S("Trust has emerged as a particularly critical factor distinguishing healthcare AI adoption from adoption in other domains. Lee and Kim (2021) examined trust dynamics in chatbot-based healthcare services, finding that healthcare emerged as the domain where trust was most carefully calibrated and where single failures had the most negative impact on subsequent usage intentions. The challenge of trust in AI systems is compounded by the often opaque nature of AI decision-making: when a chatbot responds unexpectedly, users cannot understand why the system behaved as it did, undermining trust in ways that human service failures do not."),
        H3("2.7.2 Factors Influencing Patient Engagement"),
        S("Khosravi and Azar (2024) identified clarity of purpose, consistency of behavior, and responsiveness to user emotional state as the most important predictors of sustained patient engagement with chatbot services. Spring et al. (2023) examined technical mechanisms for generating empathic responses in conversational AI, arguing that empathic responses in administrative contexts require acknowledgment of the user's emotional context, particularly when errors occur and the user expresses frustration."),
        H3("2.7.3 Trust Repair After System Failures"),
        S("Linina (2022) examined consumer satisfaction in chatbot interactions, finding that service recovery quality was one of the strongest predictors of overall satisfaction. Chatbots that responded to failure with clear, helpful error messages and recovery paths produced significantly higher satisfaction ratings than those that produced generic error notifications or fell silent. Cordero et al. (2022) similarly documented the importance of clear error communication, noting that users were generally willing to tolerate a reasonable failure rate if failures were clearly communicated and recoverable."),
        H3("2.7.4 Brand Representation and Institutional Identity"),
        S("Kühnel et al. (2020) established foundational principles about how conversational interfaces represent institutional identity, noting that chatbot interactions shape users' perceptions of the institution as much as the physical service environment does. Untari (2020) documented how government health agencies used chatbots during the COVID-19 pandemic, finding that chatbot communication quality was associated with perceptions of institutional competence and crisis management effectiveness."),
        H2("2.8 Governance and Ethics of Healthcare AI"),
        H3("2.8.1 Bioethical Foundations for Healthcare AI Governance"),
        S("Ethical Considerations in Using AI Chatbots for Culturally Sensitive Mental Health Support in African Psychotherapy (n.d.) examined the application of bioethical principles to AI systems in African healthcare contexts, arguing that standard frameworks do not fully address the ethical dimensions of AI deployment in contexts with different cultural, economic, and infrastructural characteristics. These extended ethical considerations are directly relevant to the Kenyan hospital context, where beneficence requires active improvement of access for underserved populations and justice requires attention to whether Swahili-speaking and rural users receive equivalent service quality."),
        H3("2.8.2 Data Protection and Privacy"),
        S("Calvaresi et al. (2021) demonstrated a technical architecture for privacy-compliant healthcare AI in EREBOTS, describing design patterns including data minimization, access control, purpose limitation, and audit logging that collectively enable privacy-by-design implementation. In Kenya, the Data Protection Act (2019) establishes requirements for data collection, processing, and storage that apply to all digital systems collecting personal data, requiring explicit consent from data subjects, purpose limitation for collected data, and adequate security measures."),
        H3("2.8.3 Accountability, Transparency, and Explainability"),
        S("Pears and Konstantinidis (2023) documented a progressive shift in healthcare chatbot research toward research incorporating ethical, policy, and governance dimensions. Wiens et al. (2024) emphasized that accountable AI for healthcare requires clear accountability frameworks specifying who is responsible when AI systems make consequential errors. Transparency about how AI systems make decisions is both an ethical requirement and a practical necessity for appropriate use by clinicians and administrators."),
        H3("2.8.4 Human Oversight and Escalation"),
        S("The Evolving Role of Virtual Health Assistants (n.d.) emphasized that the most successful healthcare AI deployments maintain human-in-the-loop principles not as a temporary expedient but as a permanent feature of responsible deployment. The paper argued that even highly capable AI systems should maintain escalation pathways because the range of situations encountered in real healthcare settings exceeds what can be anticipated in system design."),
        H2("2.9 Chatbots in Non-Clinical Service Settings"),
        H3("2.9.1 Customer Service Chatbots"),
        S("Auer et al. (2024) examined chatbot deployment in airport environments, finding that chatbots were most effective for routine, high-frequency queries but less effective for non-standard situations requiring contextual judgment. The design implication, that chatbots should handle common cases excellently while maintaining clear escalation pathways for edge cases, directly parallels the approach in the present thesis where the appointment booking chatbot handles standard tasks with high reliability while escalating unusual situations to human staff."),
        H3("2.9.2 Cross-Sector Evidence"),
        S("The role of chatbots in enhancing customer experience has been documented across service industries (Kumari Khushboo, 2024), with design quality identified as the primary determinant of chatbot performance: poorly designed chatbots produce worse outcomes than no chatbot at all, while well-designed chatbots consistently outperform traditional channels for covered task types. Ordemann et al. (2021) similarly found that understanding of chatbot capabilities and clear task scope were the most important predictors of user satisfaction in customer service contexts."),
        H2("2.10 Design Science Research as Methodology for Healthcare AI"),
        S("Design science research (DSR) is a research paradigm focused on the creation and evaluation of purposeful IT artifacts. Peffers et al. (2007) developed a DSR process model widely used in information systems research: identify and motivate the problem, define solution objectives, design and develop the artifact, demonstrate the artifact in realistic scenarios, evaluate the artifact against defined criteria, and communicate findings and contributions. Hevner et al. (2004) established DSR as a legitimate and rigorous research approach in information systems, articulating seven guidelines emphasizing the creation of useful artifacts, rigorous evaluation, contribution to generalizable knowledge, and communication to both academic and practice audiences."),
        H2("2.11 Synthesis of Research Gaps"),
        S("Drawing together the evidence across the reviewed literature, five significant gaps can be identified that the present thesis directly addresses. The first gap concerns integrated evaluation of multilingual reliability and transactional completion: most studies of multilingual health technology examine conversational quality or task completion separately, but few examine whether language consistency is maintained through to final transactional outputs. The second gap concerns design guidance for deterministic control within LLM assistants: practical design patterns for embedding deterministic controls within LLM conversational systems remain underdeveloped in published literature. The third gap concerns patient-facing queue analytics in conversational booking tools: the queue management literature and the chatbot literature have largely developed independently. The fourth gap concerns governance frameworks for healthcare AI in low-resource Sub-Saharan African settings. The fifth gap concerns integrated artifacts combining multiple dimensions in evaluable systems."),
        H2("2.12 Analytical Framework"),
        S("Based on the reviewed literature, this thesis adopts a five-dimension analytical framework. The accessibility dimension addresses the system's ability to serve diverse users in both English and Swahili. The reliability dimension addresses consistent and correct booking execution. The optimization dimension addresses decision support through queue-aware recommendations. The trust dimension addresses user confidence through consistent, transparent, predictable behavior. The governance dimension addresses institutional oversight and policy compliance. These five dimensions are interrelated: governance enables reliable operation, reliability supports trust, accessibility widens who benefits, and optimization provides additional value justifying institutional investment."),
        H2("2.13 Conclusion"),
        S("The reviewed literature confirms that chatbot potential in healthcare is substantial but conditional on careful attention to design, evaluation, and governance imperatives. This thesis fills identifiable gaps by presenting an integrated design science artifact that combines hybrid conversational-deterministic architecture, multilingual transactional consistency, queue-aware slot recommendation, and governance-oriented design in a single evaluable system. The five-dimension analytical framework provides a coherent lens for artifact design and evaluation grounded in the reviewed literature."),
    ]
    return title, sections


def chapter_three():
    title = "CHAPTER 3. RESEARCH METHODOLOGY"
    sections = [
        H2("3.1 Introduction"),
        S("This chapter describes the philosophical foundations, research design, methodological framework, and evaluation strategy that guide development and assessment of the AI-driven hospital appointment support artifact. The chapter explains why design science research was chosen for this study, how it was adapted for healthcare conversational AI, what design processes were followed from problem identification through artifact evaluation, and how results are analyzed and interpreted."),
        H2("3.2 Research Philosophy and Paradigm"),
        S("The philosophical orientation of this thesis is pragmatism, combined with elements of critical realism about social institutions and technologies. Pragmatism holds that knowledge should be evaluated by its consequences: whether it enables effective action in the world. In this thesis, pragmatism manifests in the explicitly practical research aim, outcome-oriented evaluation criteria, and the commitment to produce knowledge that practitioners and policy-makers can use."),
        S("Critical realism acknowledges that hospitals, patients, and administrative systems have an independent existence beyond the researcher's perspective. The problems identified in the research problem statement are real problems experienced by real patients and staff. The artifact must work in this real world, not only in a theoretically constructed research world."),
        S("The research paradigm is design science research, which occupies a distinctive methodological position by focusing on the creation of artifacts that solve real problems. Unlike positivism (which observes existing phenomena) or interpretivism (which interprets meaning in existing situations), DSR changes the world by creating new things. The rigor in DSR comes from the discipline of the design process, the evaluation against explicit criteria, and the communication of generalizable design knowledge."),
        H2("3.3 Design Science Research Framework"),
        H3("3.3.1 Foundations of Design Science Research"),
        S("Hevner et al. (2004) established design science as a legitimate and rigorous research approach in information systems, proposing seven guidelines: create a viable artifact, address a relevant problem, evaluate the artifact's utility, make contributions to knowledge, follow rigorous research methods, treat artifact design as a search process, and communicate research results to both academic and practice audiences. The present thesis follows the Peffers et al. (2007) process model adapted for healthcare AI: identify and motivate the problem, define solution objectives, design and develop the artifact, demonstrate the artifact in realistic scenarios, evaluate the artifact against defined objectives, and communicate findings and contributions."),
        H3("3.3.2 Artifact Types"),
        S("The thesis produces an instantiation artifact (a working prototype), supported by construct artifacts (the five-dimension analytical framework of accessibility, reliability, optimization, trust, and governance) and method artifacts (the hybrid conversational-deterministic architecture and the design iteration process). The instantiation provides the most direct evidence of what is possible in healthcare AI design; the constructs and methods provide generalizable knowledge extending beyond the specific prototype."),
        H3("3.3.3 Adaptation for Healthcare AI"),
        S("The adapted DSR methodology incorporates healthcare-specific requirements: clinical boundary maintenance (the artifact must clearly identify and enforce the limit between administrative support and clinical advice), patient safety considerations (evaluation must address scenarios where system failure could affect patient care access), regulatory alignment with Kenya's Data Protection Act (2019), and language equity requirements for English and Swahili. These requirements are incorporated into each phase of the DSR process from problem identification through communication of findings."),
        H2("3.4 Research Design"),
        H3("3.4.1 Overall Research Design"),
        S("The overall research design is an iterative artifact design and evaluation study following the DSR process model. The system was not designed once and then tested; instead, it evolved through multiple design cycles, with each cycle building on lessons from the previous iteration. This iterative design accommodates the inherent uncertainty of complex socio-technical problems where some requirements only become clear during implementation."),
        H3("3.4.2 Design Iteration Cycles"),
        S("Five major design iteration cycles addressed progressively more complex challenges. The first cycle established baseline conversational booking capability, examining whether LLM-based conversation was capable of understanding healthcare appointment booking requests expressed in natural language. The second cycle addressed tool-call stability and schema compliance, adding pre-invocation parameter validation and retry logic. The third cycle introduced a deterministic state machine defining valid workflow states and mandatory information prerequisites. The fourth cycle implemented language support and deterministic localization for transaction-critical outputs. The fifth cycle integrated all components, tested complex scenarios, and documented governance implications."),
        H3("3.4.3 Validation Gates Between Cycles"),
        S("Each cycle concluded with a validation phase checking: no regressions in prior functionality, new functionality working reliably in dedicated test scenarios, code quality standards maintained, and documentation adequate for future maintenance. Only when all validation checks passed did the next cycle begin."),
        H2("3.5 Case and Setting"),
        S("The research case is the administrative scheduling challenge at Kenyatta University Teaching, Referral and Research Hospital, used as a representative context for Kenyan referral hospital scheduling needs. KUTRRH is a 500-bed tertiary teaching and referral facility serving diverse patient populations. The evaluation is conducted as a controlled prototype evaluation rather than a live deployment study, allowing rigorous testing of edge cases and error conditions without risk to real patient care."),
        H2("3.6 Artifact Description and Architecture"),
        S("The artifact consists of seven integrated subsystems: the frontend conversational interface (Streamlit-based web chat), the orchestration graph (LangGraph-based workflow state machine), the LLM-powered language understanding component (Claude via Anthropic Python SDK), the deterministic parser and guardrail module (rule-based validation), the appointment operation tools (seven operations with strict schemas), the queue estimation and recommendation subsystem (deterministic scoring model combining service-type baseline, time-of-day, and day-of-week factors), and the logging and audit subsystem (structured event logging)."),
        H2("3.7 Data Sources and Management"),
        S("Design and implementation data are produced by the researcher through the design process and include source code, design documents, and evaluation records. Synthetic test data reflects realistic patient and appointment patterns without using real patient clinical information, enabling systematic construction of test cases that exercise specific system behaviors including error conditions. Secondary evidence from the literature review informs design choices and contextualizes evaluation findings. No real patient clinical data is used at any point in the research."),
        H2("3.8 Evaluation Strategy"),
        H3("3.8.1 Evaluation Methods"),
        S("The evaluation combines three primary methods. Functional testing verifies that each defined system function activates correctly and produces correct outputs in standard cases. Scenario-based evaluation uses a library of 47 scenarios across nine categories: standard English booking flows, standard Swahili booking flows, mixed-language flows, partial information flows, invalid input flows, error recovery flows, cancellation flows, next-available lookup flows, and queue recommendation flows. Governance assessment evaluates the artifact's compliance with a defined set of governance requirements derived from the literature review and Kenya's Data Protection Act."),
        H3("3.8.2 Metrics and Criteria"),
        S("For reliability evaluation, primary metrics are booking completion rate, invalid transaction prevention rate, and error recovery rate. For multilingual consistency, the primary metric is language consistency rate in transaction-critical outputs. For queue recommendations, metrics include recommendation activation rate, interpretability assessment, and preference uptake. For governance, metrics include audit log completeness, data minimization compliance, scope adherence, and escalation routing performance."),
        H2("3.9 Reliability, Validity, and Research Quality"),
        S("Research reliability is enhanced through deterministic system components, precisely documented test scenarios, and consistent evaluation on the same system implementation. Internal validity is supported through iterative testing with explicit documentation of what changed between iterations. External validity is addressed through documentation of design principles at a level of abstraction enabling application beyond this specific implementation. Research ethics are addressed through the use of synthetic evaluation data and documentation of the governance framework for deployment. No human participants were involved in the evaluation."),
        H2("3.10 Limitations of the Methodology"),
        S("Key limitations include prototype-scale evaluation without production load testing; researcher-conducted evaluation with potential confirmation bias; evaluation by researchers rather than actual patients, meaning user acceptance is not established; non-native-speaker Swahili evaluation that may miss naturalness issues; and queue prediction calibrated to general rather than local hospital patterns. These limitations do not invalidate findings but contextualize them as evidence of technical feasibility and design patterns rather than proof of real-world effectiveness."),
    ]
    return title, sections


def chapter_four():
    title = "CHAPTER 4. SYSTEM DESIGN AND IMPLEMENTATION"
    sections = [
        H2("4.1 Introduction"),
        S("This chapter details the architecture and implementation of the AI-driven hospital patient support assistant, presenting the design decisions that address the research objectives and the specific mechanisms through which the system achieves reliable booking completion, queue-aware recommendation, and multilingual localization consistency. The chapter is organized to follow the layered architecture of the system, beginning with the principles that guided all design decisions and proceeding through each architectural layer and key subsystem."),
        H2("4.2 Design Philosophy and Foundational Principles"),
        S("Five foundational principles guided all architectural and implementation decisions. The first principle is separation of language understanding from transaction execution: the LLM handles interpretation and generation, while deterministic logic validates all parameters before any transaction execution. This architectural boundary prevents conversational errors from propagating to transactional errors."),
        S("The second principle is deterministic control for high-risk workflow transitions: booking transitions require verification that all mandatory parameters are present, valid, and consistent before the booking tool is invoked. The third principle is minimal but sufficient data capture: the system collects only patient name, patient ID, phone, email, service type, appointment date, and appointment time."),
        S("The fourth principle is explainable recommendation outputs: queue recommendations include predicted congestion level, estimated wait time range, and brief plain-language explanation enabling patients to evaluate recommendations against their own priorities. The fifth principle is policy-aware logging and error handling: all significant events are logged in a structured format supporting both debugging and auditing, and error handling is user-facing and recovery-oriented."),
        H2("4.3 System Architecture"),
        H3("4.3.1 Layered Architecture Design"),
        S("The system employs a five-layer architecture that supports independent development, testing, and modification of each layer. Layer 1 is the Interface Layer: how users interact with the system (Streamlit-based web chat). Layer 2 is the Orchestration Layer: how the system sequences operations (LangGraph state machine). Layer 3 is the Intelligence Layer: how the system understands language and generates responses (Claude LLM via Anthropic Python SDK). Layer 4 is the Tool and Operations Layer: how the system affects appointment records (seven operations with strict schemas). Layer 5 is the Data and Logging Layer: how the system persists information and creates audit trails (JSON appointment store and structured event logging)."),
        H3("4.3.2 Interface Layer Implementation"),
        S("The patient-facing interface is implemented as a web-based chat application using the Streamlit framework. The design follows a familiar messaging application pattern, maintaining and displaying the full conversation history within a session and including a visible scope disclaimer at the start of each session. Session state management persists conversation history, detected language context, partially-collected booking information, and user preferences across messages within a session."),
        H3("4.3.3 Orchestration Layer Implementation"),
        S("The orchestration layer is built on the LangGraph framework, implementing the booking workflow as a directed graph where nodes represent states and edges represent conditional transitions. Principal states include initial, intent-detected, collecting-patient-details, service-confirmed, date-selected, time-selected, pre-booking-validation, booking-executing, booking-confirmed, and booking-failed. State transitions are conditional on information completeness and validity. The workflow also includes alternative paths for cancellation and information retrieval with their own validation requirements."),
        H3("4.3.4 Intelligence Layer Implementation"),
        S("The intelligence layer uses Claude (Anthropic) accessed through the Anthropic Python SDK. The system prompt establishes the assistant's identity, scope, and operating principles, and is updated dynamically based on the current workflow state to focus the LLM's attention on the most relevant information at each step. Tool definitions are provided in Anthropic's tool use format. Tool invocations suggested by the LLM are validated by the orchestration layer before execution."),
        H3("4.3.5 Tool and Operations Layer Implementation"),
        S("Seven tools implement appointment operations: create-appointment, cancel-appointment, get-next-available, recommend-best-slot, check-availability, get-appointment-details, and list-services. The create-appointment tool requires patient name, ID, phone, email, service type, date, and time as mandatory parameters, and validates all inputs before creating a record. On success, it returns a structured confirmation including appointment ID, confirmed date and time, service name, location, and patient instructions."),
        H3("4.3.6 Data and Logging Layer Implementation"),
        S("Appointment records are stored in JSON format for the prototype, with each record including a unique appointment ID, patient details, service type, date and time, appointment status, creation timestamp, and modification timestamps. The logging subsystem writes structured log records for all significant system events using a consistent schema: timestamp, session ID, event type, event data, and outcome."),
        H2("4.4 Booking Workflow Design in Detail"),
        H3("4.4.1 Information Collection Sequence"),
        S("The booking workflow accepts information in whatever order the user provides it while tracking what has been collected. Patient identification information (name, ID, phone, email) is collected first; appointment specification (service, date, time) is collected second. The state machine tolerates information provided across multiple turns and only prompts for information that has not yet been collected."),
        H3("4.4.2 Service Type Resolution"),
        S("Service type resolution uses a two-stage approach: LLM initial matching of the user's expression to a defined service type, followed by explicit user confirmation before proceeding. A bilingual service type lookup table maintains the authorized mapping between natural language expressions in both English and Swahili and service system identifiers, ensuring consistent mapping regardless of how the LLM interprets a specific expression."),
        H3("4.4.3 Date and Time Parsing"),
        S("Date and time parsing uses LLM interpretation for natural language expressions combined with deterministic validation of resulting values. Dates must be in the future; times must be within clinic hours (8:00 AM to 5:00 PM); the service must be available on the requested day. Ambiguous date references trigger explicit disambiguation questions to the user to prevent booking errors caused by misinterpreted relative date references."),
        H2("4.5 Multilingual Localization Engine"),
        H3("4.5.1 Language Detection and Context Tracking"),
        S("Language detection operates at the conversation level using a combination of LLM-based language identification and simple string matching for common function words. A single message in a different language does not trigger a language context switch; sustained use across two or more turns does. The detected language context is stored in session state and passed to all subsequent system components, ensuring consistency across the entire interaction."),
        H3("4.5.2 Deterministic Localization of Transaction-Critical Outputs"),
        S("Booking confirmations, error messages, and queue recommendation summaries are generated through deterministic localization templates rather than LLM generation. For each message type, two templates exist (English and Swahili), with the selection determined by session language context. Service type labels and appointment type labels are maintained in bilingual lookup tables. This architectural decision addresses the language drift problem by removing language-consistency decisions from the probabilistic language model for the most critical messages."),
        H3("4.5.3 Conversational Language Consistency"),
        S("Conversational responses rely on LLM language consistency guided by explicit system prompt instruction. This approach is sufficient for conversational exchanges where some variation is acceptable, but is not relied upon for transaction-critical outputs where exact consistency is required."),
        H2("4.6 Queue Prediction and Recommendation Subsystem"),
        H3("4.6.1 Congestion Prediction Model"),
        S("The queue prediction model is a deterministic scoring function combining three factors: service-type baseline congestion (reflecting different demand levels across services), time-of-day factor (reflecting consistent patterns where mid-morning is busiest and early afternoon is least congested), and day-of-week factor (reflecting Monday peak and mid-week evenness). The combined score is normalized to a zero-to-one scale and mapped to qualitative categories (low, moderate, high congestion) with associated estimated wait time ranges."),
        H3("4.6.2 Recommendation Generation and Presentation"),
        S("Recommendations include the slot date and time (in user-appropriate format for the detected language), the congestion level label (Low/Moderate/High in English or Chini/Wastani/Juu in Swahili), an estimated wait time range, and a brief explanation. Recommendations are presented as a ranked list with the lowest-congestion option highlighted, and users can select any presented option or request alternatives."),
        H3("4.6.3 Handling Prediction Uncertainty"),
        S("Uncertainty is communicated through disclaimer language included in every recommendation presentation: 'These estimates are based on typical patterns. Actual waiting times may vary based on clinic conditions on the day of your appointment.' This transparent communication maintains user trust even when actual waits differ from estimates."),
        H2("4.7 Governance Controls Implementation"),
        H3("4.7.1 Scope Enforcement"),
        S("Scope enforcement uses system prompt instructions and intent classification to identify clinical requests and redirect them appropriately. When a clinical request is detected, the system acknowledges that it cannot help with clinical questions, explains that it is an appointment scheduling assistant, and offers to help the user book an appointment with the appropriate specialist. This response pattern maintains helpful engagement without overstepping the administrative scope."),
        H3("4.7.2 Escalation Pathways"),
        S("Escalation is triggered by three consecutive booking failures, explicit requests for human assistance, detection of emergency language, and technical errors preventing normal operation. Escalation messages are generated in the user's language from deterministic templates and include the hospital's contact information and guidance on next steps."),
        H3("4.7.3 Audit Logging and Data Governance"),
        S("The audit system captures all significant events in structured log records: message receipt, intent classification, entity extraction, tool invocations and results, state transitions, and session start and end. Automatic session data clearing at session end minimizes data persistence beyond the immediate interaction. The structured log format enables automated analysis and supports the audit requirements that healthcare governance demands."),
        H2("4.8 Chapter Summary"),
        S("This chapter has detailed the design and implementation of the AI-driven hospital appointment support artifact across five architectural layers. The design is guided by five foundational principles that translate literature insights into concrete design commitments. The layered architecture separates concerns between interface, orchestration, intelligence, tools, and data, enabling independent development and testing of each layer. The multilingual localization engine combines LLM-based language understanding for conversational exchanges with deterministic template-based generation for transaction-critical outputs, addressing the language drift problem through architectural design rather than model instruction."),
    ]
    return title, sections


def chapter_five():
    title = "CHAPTER 5. RESULTS AND EVALUATION"
    sections = [
        H2("5.1 Introduction"),
        S("This chapter presents evaluation findings from systematic functional and scenario-based testing of the AI-driven hospital appointment support artifact. The evaluation was structured to assess performance across five analytical dimensions: accessibility, reliability, optimization, trust, and governance. Within each dimension, specific metrics were applied through a library of 47 controlled test scenarios distributed across nine categories."),
        H2("5.2 Evaluation Context and Testing Environment"),
        S("All evaluation was conducted on the integrated prototype running on a researcher workstation. Test data was synthetic: patient names, identification numbers, contact details, and appointment histories were fabricated but realistic in structure. Each scenario was executed at least twice to verify result consistency. The 47 scenarios were distributed as follows: standard English booking flows (9 scenarios), standard Swahili booking flows (7 scenarios), mixed-language flows (5 scenarios), partial information flows (8 scenarios), invalid input flows (7 scenarios), error recovery flows (5 scenarios), cancellation flows (4 scenarios), next-available lookup flows (3 scenarios), and queue recommendation flows (7 scenarios)."),
        H2("5.3 Functional Coverage Evaluation"),
        S("All seven core functions activated successfully across all applicable scenarios. Appointment booking activated in all 24 applicable scenarios, optimal slot recommendation in all 7 recommendation scenarios, wait time estimation was internally consistent across all 7 recommendation scenarios, alternative slot suggestion activated proactively in all 5 high-congestion scenarios, appointment cancellation activated in all 4 cancellation scenarios, next-available retrieval activated in all 3 retrieval scenarios, and busy and quiet time identification activated correctly in all 5 applicable scenarios. Overall functional activation rate was 100%."),
        H2("5.4 Reliability Evaluation: Booking Completion"),
        H3("5.4.1 Baseline Completion Rate"),
        S("Pre-guardrail baseline testing across 20 scenarios produced 17 successful bookings, yielding an 85% completion rate. The three failures occurred in one scenario where date was provided before service type (confusing the LLM's state tracking), and two scenarios where the LLM generated malformed tool invocations with missing required parameters."),
        H3("5.4.2 Post-Guardrail Completion Rate"),
        S("Post-guardrail testing across the same 20 scenarios produced 20 successful bookings, yielding a 100% completion rate. The previously failing scenarios succeeded because the state machine correctly held partial information and continued collecting missing data, and the guardrail validation caught missing parameters before tool invocation and generated appropriate clarifying prompts. This improvement from 85% to 100% represents elimination of a systematic failure class rather than incremental performance improvement."),
        H3("5.4.3 Invalid Transaction Prevention"),
        S("All seven invalid input scenarios were caught before tool invocation and before reaching the appointment data store: booking without patient ID, past date booking, unrecognized service type, outside-clinic-hours booking, invalid phone number format, conflicting slot booking, and missing email address. In every case, the guardrail validation identified the invalid condition and generated a specific, actionable error message explaining what was required to proceed."),
        H3("5.4.4 Non-Standard Input Handling"),
        S("All eight non-standard input scenarios succeeded, including providing appointment details before patient details, using relative time expressions such as 'the day after tomorrow at lunchtime,' providing partial information across multiple short messages, using implicit service references, using colloquial service expressions, and providing service descriptions rather than service names. Three of the eight scenarios required at least one clarifying question from the system before all information was resolved, and in all three cases the question was appropriately specific."),
        H2("5.5 Multilingual Consistency Evaluation"),
        H3("5.5.1 English-Only Interactions"),
        S("Language consistency in nine English-only scenarios was 100%. All conversational responses, transaction-critical outputs including confirmations, error messages, and recommendations, and all user-facing text was in English throughout each scenario. No language drift was observed in pure-English interactions."),
        H3("5.5.2 Swahili-Only Interactions"),
        S("Conversational response language consistency in seven Swahili scenarios was 100%: the LLM consistently responded in Swahili as instructed. Transaction-critical output consistency was also 100% using deterministic localization templates. Pre-deterministic testing showed language drift in transaction-critical outputs in approximately 43% of Swahili scenarios, confirming that the deterministic localization architecture fully resolves the language drift problem."),
        H3("5.5.3 Mixed-Language Interaction Handling"),
        S("Language context detection correctly identified the dominant language in all five mixed-language scenarios. The system produced responses consistent with the identified dominant language in all cases, including final booking confirmations. The mid-conversation language switch scenario produced the expected behavior: the system continued in the initial language for two turns after the switch, then updated the language context after detecting persistent use of the new language."),
        H3("5.5.4 Pre-Deterministic versus Post-Deterministic Localization Comparison"),
        S("Pre-deterministic consistency across 12 Swahili and mixed-language scenarios was 57% (7 of 12 transaction-critical outputs in the expected language). Post-deterministic consistency across the same 12 scenarios was 100% (12 of 12 correct). This improvement directly demonstrates the value of the deterministic localization architecture for multilingual healthcare applications."),
        H2("5.6 Queue Recommendation Evaluation"),
        S("Recommendations activated in all 7 applicable scenarios with consistent presentation format including congestion label, wait time estimate, explanation, and ranked options. Interpretability assessment rated all 7 scenarios positive on all four criteria: congestion label clarity, wait time specificity, explanation appropriateness, and overall actionability. In 6 of 7 scenarios (86%), the scripted user selected the recommended lowest-congestion slot when it was reasonably convenient. The uncertainty disclaimer appeared in all 7 scenarios and was assessed as appropriately toned."),
        H2("5.7 Error Recovery Evaluation"),
        S("All five error recovery scenarios produced appropriate system behavior with clear recovery paths. The service unavailable scenario produced graceful degradation with alternative contact information and a continued booking offer. The database connection error scenario produced a retry offer with backup contact. The date conflict scenario produced an apologetic redirect to available alternatives with their congestion ratings. The invalid appointment ID scenario produced a verification suggestion. The session timeout scenario produced a summary of prior session data and a continuation offer. In no scenario was the user left without a clear path forward."),
        H2("5.8 Governance Evaluation"),
        S("Audit log completeness was 100% across all 10 governance-evaluated scenarios: all defined event types (message-received, intent-classified, entity-extracted, state-transition, tool-invoked, tool-completed, error-occurred, session-ended) were captured with appropriate data including timestamps and session IDs. Data minimization compliance was 100%: only the five defined required fields were collected in any scenario and no additional personal data was solicited. Scope adherence was 100%: all three out-of-scope requests (medication dosage, symptom assessment, specialist recommendation) were correctly redirected with appropriate acknowledgment and administrative assistance offers. Escalation routing activated correctly in both escalation test scenarios."),
        H2("5.9 Synthesis of Evaluation Findings"),
        S("The following summary metrics characterize the artifact's performance across all evaluation dimensions. Booking completion rate: 100% post-guardrail versus 85% pre-guardrail baseline. Invalid transaction prevention rate: 100% across seven invalid input scenarios. Error recovery path availability: 100% across five error scenarios. Multilingual consistency in transaction-critical outputs: 100% post-deterministic localization versus 57% pre-deterministic. Functional activation rate: 100% across all seven functions and all applicable scenarios. Recommendation interpretability: 100% positive assessment across all four interpretability criteria. Audit log completeness: 100%. Data minimization compliance: 100%. Scope adherence: 100%."),
        S("Qualitative observation noted that conversational quality was consistently appropriate throughout all scenarios, with natural and contextually appropriate response language. Swahili interactions were functional but somewhat more formal than natural spoken Swahili, a consequence of the LLM's training data distribution. Queue recommendation presentations were clear, professional, and evidence-based in character."),
    ]
    return title, sections


def chapter_six():
    title = "CHAPTER 6. DISCUSSION"
    sections = [
        H2("6.1 Introduction"),
        S("This chapter interprets evaluation findings in relation to the five research questions, the literature reviewed in Chapter 2, the specific context of Kenyan healthcare administration, and broader implications for conversational healthcare AI design and deployment. The discussion moves progressively from specific findings to general principles, situating evidence from this thesis within the wider conversation about healthcare AI design, multilingual digital health equity, and AI governance."),
        H2("6.2 Research Question 1: Architecture for Reliable Healthcare AI"),
        S("Research Question 1 asked what modular architecture best combines conversational flexibility with reliable hospital transaction execution and what design principles should govern boundaries between components. The evaluation confirms that the five-layer architecture with explicit separation of concerns is effective for healthcare administrative AI. The pre-guardrail to post-guardrail improvement (85% to 100%) is not marginal but represents elimination of a systematic failure class. This aligns with the finding of Caldarini et al. (2022) that reliable commercial chatbot deployments combine neural language understanding with deterministic action execution."),
        S("The architectural contribution extends to institutional adoption: the modular architecture enables incremental adoption, component upgrading without system-wide changes, and independent testing of each layer. These properties support long-term sustainability in a rapidly evolving AI technology landscape. The specific boundary design principles that proved effective are: the LLM suggests, deterministic logic validates; conversational errors produce clarifying questions, not transaction failures; and language understanding is separated from language output in critical messages."),
        H2("6.3 Research Question 2: Deterministic Guardrails and Reliability"),
        S("Research Question 2 asked how deterministic guardrail mechanisms affect booking completion quality, error recovery behavior, and user experience. Deterministic guardrails materially improved booking reliability across multiple dimensions. The post-guardrail failures detected in pre-guardrail testing were systematic: the same categories of input produced failures on repeated runs. Post-guardrail, these categories are structurally impossible failures, not merely reduced-probability failures. This finding challenges the binary framing in earlier chatbot literature that positioned rule-based and generative systems as opposites (Ordemann et al., 2021), demonstrating that the two approaches are complementary rather than competing in the specific context of healthcare transactional AI."),
        H2("6.4 Research Question 3: Queue-Aware Recommendations"),
        S("Research Question 3 asked how predicted congestion indicators and interpretable queue recommendations can be integrated into conversational booking workflows. The evaluation found that the integration is feasible and produces interpretable outputs that influence user decisions. The key design choices enabling this outcome were plain-language presentation, honest uncertainty disclosure, brief explanatory context, and natural integration into the booking conversation. The 86% uptake of recommended slots in scripted testing suggests potential for meaningful demand redistribution at population scale. Linina (2022) cautioned that queue information must be calibrated accurately to be trusted: production deployment would require calibration from actual hospital data to achieve trustworthy prediction accuracy."),
        H2("6.5 Research Question 4: Multilingual Localization Consistency"),
        S("Research Question 4 asked what technical and design approaches ensure consistent language localization in transactional messages throughout the full booking workflow. The evaluation provides unambiguous evidence: improvement from 57% to 100% consistency using deterministic localization demonstrates conclusively that LLM-based language consistency is insufficient for transaction-critical messages and that deterministic localization completely addresses this gap. The underlying reason is structural: probabilistic language models optimize for plausibility, not for consistency. Deterministic templates guarantee consistency regardless of conversational context."),
        S("This finding has implications beyond the Kenyan English-Swahili context. Any domain where transaction-critical messages must be delivered consistently in user-preferred languages would benefit from this architectural approach. The finding extends the contribution of Panagiotidis (2024), who documented LLM multilingual capabilities, by identifying the specific boundary where LLM-based consistency is insufficient and deterministic approaches are required."),
        H2("6.6 Research Question 5: Governance and Policy Alignment"),
        S("Research Question 5 asked what ethical safeguards, accountability mechanisms, audit capabilities, and governance controls are necessary for responsible deployment in Kenyan hospital contexts. The governance evaluation found that built-in controls address primary governance requirements identified in the literature and in Kenya's Data Protection Act (2019). The deeper insight is that governance controls are most effective when built into the architecture rather than applied as external oversight mechanisms. This 'governance by design' approach, aligned with the recommendations of Calvaresi et al. (2021) in their EREBOTS platform, enables faster institutional adoption by providing evidence of responsible design before deployment."),
        H2("6.7 Broader Implications for Healthcare AI Design"),
        H3("6.7.1 Hybrid Architecture as a Standard Approach"),
        S("The collective evidence makes a strong case that hybrid conversational-deterministic architecture should become the standard approach for healthcare administrative AI, replacing both purely rule-based and purely generative alternatives. The specific evidence base includes: better reliability outcomes than purely generative approaches, better linguistic flexibility than purely rule-based approaches, and compatibility with trust-building, error-recovery, and governance requirements that healthcare deployment demands."),
        H3("6.7.2 Implications for Digital Health Equity"),
        S("The multilingual localization findings have direct implications for digital health equity. The deterministic localization approach demonstrated in this thesis is a concrete technical mechanism for achieving language equity in conversational health services. Institutions implementing this approach can provide genuinely equivalent service quality to speakers of all supported languages, not merely cosmetic multilingual support. As noted by Skjuve and Brandtzæg (2018), language accessibility is not merely a convenience but a determinant of whether a service is genuinely accessible to diverse populations."),
        H3("6.7.3 Implications for AI Governance in Low-Resource Settings"),
        S("The governance framework documented in this thesis was developed for the Kenyan context but draws on principles with wider applicability. The finding that governance by design is more effective than governance by oversight alone applies in any institutional context. For low-resource settings specifically, automated governance controls reduce the institutional compliance burden by collecting evidence for review automatically, without requiring large compliance staff for routine operation."),
        H2("6.8 Study Limitations and Their Implications"),
        S("The most significant limitation is that evaluation was conducted with scripted scenarios and synthetic data rather than real patients and natural interaction. The controlled evaluation is best understood as evidence of technical capability rather than evidence of real-world effectiveness with actual patients. Pilot deployment with real patients is the essential next step for establishing real-world effectiveness."),
        S("The researcher-conducted evaluation introduces the risk of inadvertent bias. The Swahili evaluation was conducted by a non-native speaker who can assess functional accuracy but may not assess naturalness with the sensitivity of a native speaker. The queue prediction model is calibrated to general rather than local hospital patterns. These limitations are explicitly acknowledged to contextualize the strength of evidence for the findings."),
    ]
    return title, sections


def chapter_seven():
    title = "CHAPTER 7. CONCLUSION AND RECOMMENDATIONS"
    sections = [
        H2("7.1 Synthesis of Contributions"),
        S("This thesis set out to design, implement, and evaluate a multilingual AI-driven appointment support assistant for Kenyan hospital contexts using a design science approach and a pragmatic research philosophy. The research has produced four major contributions: a primary artifact contribution (a working, evaluated prototype demonstrating 100% booking completion reliability, 100% transaction-critical language consistency, and governance-ready architecture), a design knowledge contribution (transferable architecture patterns and design principles), an analytical framework contribution (the five-dimension framework of accessibility, reliability, optimization, trust, and governance), and a governance framework contribution (a practical starting point for healthcare AI governance policies)."),
        H2("7.2 Key Findings Revisited"),
        S("The first key finding is that hybrid conversational-deterministic architecture is necessary and sufficient for reliable healthcare administrative AI. The evidence is: pre-guardrail evaluation demonstrating 15% booking failure rate in controlled testing, and post-guardrail evaluation demonstrating 0% failure rate with no new categories of failure introduced. The necessity is established by the evidence of LLM-only failures; the sufficiency is established by the 100% completion rate after guardrail implementation."),
        S("The second key finding is that language consistency in transaction-critical outputs requires deterministic implementation. LLM-generated Swahili confirmations produced language drift in 43% of applicable scenarios; deterministic templates produced 100% consistency in the same scenarios. This is a structural finding, not a model quality finding: the probabilistic nature of language models is incompatible with the deterministic requirement for consistency in high-stakes messages."),
        S("The third key finding is that queue recommendations improve decision quality when presented interpretably with honest uncertainty disclosure. The 86% uptake of recommended slots in scripted testing suggests meaningful demand redistribution potential. The fourth key finding is that governance by design is achievable and enables faster institutional adoption by providing evidence of responsible design before deployment."),
        H2("7.3 Implementation Recommendations for Healthcare Institutions"),
        H3("7.3.1 Readiness Assessment Before Deployment"),
        S("Before deploying a conversational AI appointment system, healthcare institutions should assess their readiness across five dimensions: technical readiness (stable appointment data infrastructure, adequate server capacity, IT staff capable of deploying and maintaining the system), governance readiness (institutional AI oversight policies, designated accountability, established complaint processes), data readiness (consent mechanisms, data protection framework aligned with Kenya's Data Protection Act, secure data storage), staff readiness (training on system capabilities and escalation handling), and patient communication readiness (clear patient-facing materials explaining chatbot functions and limitations)."),
        H3("7.3.2 Pilot Deployment Strategy"),
        S("Institutions should deploy initially as a limited pilot in a single high-volume department with strong administrative staff engagement. The pilot should run for a minimum of three to six months, monitoring booking completion rates, error rates, escalation frequencies, language distribution of interactions, and patient feedback. Pilot evaluation should compare performance against pre-deployment baselines. A booking completion rate of at least 95% and an escalation rate below 10% should be targets before expanding to additional departments."),
        H3("7.3.3 Integration with Hospital Information Systems"),
        S("Production deployment requires integration with the hospital's actual appointment and patient management systems. Key integration points are the appointment creation and cancellation interfaces, the availability query interface, the patient record verification interface, and the notification system interface for confirmation messages. Each integration point should be implemented with appropriate error handling that directs users to alternatives rather than failing silently."),
        H3("7.3.4 Queue Model Calibration"),
        S("The queue prediction model should be calibrated to local hospital data before deployment using historical appointment volume data by service type, time of day, and day of week, along with actual patient wait time records. A healthcare data analyst working with operations staff can develop a locally calibrated model. Periodic recalibration every six to twelve months accounts for systematic changes in patient volume patterns."),
        H2("7.4 Policy Recommendations"),
        H3("7.4.1 Recommendations for Kenya Ministry of Health"),
        S("The Kenya Ministry of Health should consider a national framework for healthcare AI governance clarifying regulatory expectations for AI systems in health administration. The framework should establish minimum standards for data protection, audit logging, human oversight, and patient communication, reducing the uncertainty that currently inhibits institutional AI adoption. Language equity standards should establish that digital health services serving multilingual populations must provide equivalent service quality in all officially supported languages, specifically requiring that transaction-critical outputs be provided in the patient's preferred language."),
        S("Mandatory incident reporting requirements for serious AI system failures would create visibility into real-world performance and enable learning across the healthcare system. Funding for multilingual AI development for Kenyan language contexts would accelerate the availability of high-quality AI tools in Swahili and other Kenyan languages."),
        H3("7.4.2 Recommendations for Healthcare Regulatory Authorities"),
        S("Healthcare regulatory authorities should update their frameworks for health technology evaluation to include AI-specific guidance on dynamic systems, requirements for human oversight mechanisms, and standards for AI system audit trails in healthcare. Patient data governance for AI systems should be specifically addressed in healthcare AI regulations, going beyond the Data Protection Act's general provisions to address the specific characteristics of AI systems including conversation log retention and deletion policies."),
        H3("7.4.3 Recommendations for Educational Institutions"),
        S("Graduate programs in health informatics, computer science, and ICT policy should include curriculum covering the design and governance of healthcare AI systems. Research programs should prioritize controlled pilot studies of healthcare AI systems in Kenyan and comparable African contexts, funding the transition from proof-of-concept prototypes to evidence-based real-world deployments."),
        H2("7.5 Directions for Future Research"),
        H3("7.5.1 Real-World Pilot Studies"),
        S("The most important direction for future research is controlled pilot deployment of the appointment support system in an actual hospital setting with real patients. A rigorous pilot study design would include pre-deployment baseline measurement of existing administrative data, systematic data collection during deployment, post-deployment analysis comparing outcomes against baseline, and qualitative investigation of patient and staff experiences through interviews or focus groups."),
        H3("7.5.2 Extended Language Support"),
        S("The current artifact supports English and Swahili. Kenya's linguistic diversity extends to numerous regional languages including Kikuyu, Luo, Kamba, Kalenjin, and others. Research on extending language support should examine the feasibility of LLM-based understanding of Kenyan regional languages, the requirements for native-speaker input in developing localization templates, and the interaction patterns specific to users of each language."),
        H3("7.5.3 Integration with Clinical Workflows"),
        S("Future research could examine how the administrative appointment support system interfaces with clinical functions: pre-appointment instructions, appointment reminders that reduce no-show rates, and post-appointment follow-up communications. Each extension must maintain the critical boundary between administrative support and clinical advice, and each would require appropriate governance review."),
        H3("7.5.4 Economic Impact Assessment"),
        S("An important dimension not addressed in this thesis is the economic impact of the appointment support system. Institutional savings might include reduced administrative staff time, reduced no-show rates, improved resource utilization from better demand distribution, and reduced cost per booking transaction. Patient savings might include reduced travel costs and reduced time cost for efficient online booking. Quantifying these impacts would strengthen the business case for institutional investment and inform policy decisions about digital health funding priorities."),
        H3("7.5.5 Fairness and Equity Analysis"),
        S("A systematic fairness and equity analysis of the system's performance across demographic groups would address concerns about AI systems potentially reproducing or amplifying existing healthcare access inequities. Specific questions to investigate include: whether elderly users achieve comparable booking completion rates to younger users, whether users with lower digital literacy experience higher error rates, whether rural users with poor connectivity face additional barriers, and whether the system's Swahili support serves all Swahili speaker demographics equivalently."),
        H2("7.6 Final Conclusion"),
        S("This thesis has established that hybrid conversational-deterministic architecture for AI-driven hospital appointment support is technically feasible, operationally effective, and governance-ready for Kenyan healthcare contexts. The five research questions have been comprehensively addressed: architecture design that separates language understanding from transaction execution; deterministic guardrails that raise booking completion from 85% to 100%; queue-aware recommendations that achieved 86% uptake in testing and 100% interpretability assessment; deterministic localization that raises transaction-critical output language consistency from 57% to 100%; and governance controls that satisfy all audit, minimization, scope, and escalation requirements."),
        S("Healthcare AI that is both useful and trustworthy is achievable. This thesis demonstrates how, provides the blueprint for others to follow, and establishes a clear research agenda for the evidence needed to move from prototype to policy-supported, institution-ready deployment. The patients who will benefit most from a well-designed appointment assistant are precisely those for whom the current administrative layer of healthcare is most difficult to navigate. That these patients can be served equitably and reliably with current technology is both an opportunity and an obligation for healthcare institutions and policy-makers in Kenya and comparable contexts."),
    ]
    return title, sections


# ── MAIN ──────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("Building thesis Word document ...")
    doc = build_doc()
    output_path = os.path.join(
        os.path.dirname(__file__),
        "AI_Patient_Support_Thesis_Kung'u_Kelvin_Mathigi_2026.docx"
    )
    doc.save(output_path)
    print(f"Saved: {output_path}")
