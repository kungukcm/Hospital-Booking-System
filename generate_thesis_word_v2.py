"""
Master generator — Expanded Thesis Word Document
Imports all content modules and generates the expanded thesis.
Run: python generate_thesis_word_v2.py
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

from thesis_content_ch1 import chapter_one
from thesis_content_ch2 import chapter_two
from thesis_content_ch3456 import chapter_three, chapter_four, chapter_five, chapter_six
from thesis_content_ch7 import chapter_seven


# ── helpers ──────────────────────────────────────────────────────────────────

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(2.54)


def setup_styles(doc):
    styles = doc.styles

    # Normal
    n = styles['Normal']
    n.font.name = 'Times New Roman'
    n.font.size = Pt(12)
    n.paragraph_format.space_after = Pt(0)
    n.paragraph_format.line_spacing = Pt(24)   # double-spaced
    n.paragraph_format.first_line_indent = Inches(0.5)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)
    h1.paragraph_format.first_line_indent = Inches(0)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(6)
    h2.paragraph_format.first_line_indent = Inches(0)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.font.italic = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after  = Pt(6)
    h3.paragraph_format.first_line_indent = Inches(0)


def body(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    return p


def heading(doc, text, level):
    style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}.get(level, 'Normal')
    p = doc.add_paragraph(text, style=style)
    return p


def ref_entry(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent       = Inches(0.5)
    p.paragraph_format.space_after       = Pt(0)
    p.paragraph_format.line_spacing      = Pt(24)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_page_break(doc):
    doc.add_page_break()


# ── front matter ─────────────────────────────────────────────────────────────

def front_matter(doc):
    # Title page
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.first_line_indent = Inches(0)
    r = t.add_run("HANDONG GLOBAL UNIVERSITY")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.first_line_indent = Inches(0)
    r2 = t2.add_run("Graduate School of Global Management and Information Systems")
    r2.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.paragraph_format.first_line_indent = Inches(0)
    r3 = t3.add_run(
        "AI-Driven Multilingual Appointment Support for Kenyan Hospitals:\n"
        "A Design Science Approach"
    )
    r3.bold = True
    r3.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph()

    for line in [
        "A Thesis Submitted in Partial Fulfillment",
        "of the Requirements for the Degree of",
        "Master of Science in ICT Policy and Management",
        "",
        "by",
        "",
        "Kung'u Kelvin Mathigi",
        "",
        "Student ID: [Student ID]",
        "",
        "Supervisor: [Supervisor Name]",
        "",
        "2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.add_run(line).font.size = Pt(12)

    add_page_break(doc)

    # Abstract
    heading(doc, "ABSTRACT", 1)
    body(doc,
        "Healthcare systems in Kenya and across sub-Saharan Africa face a dual administrative "
        "burden: patients struggle to book appointments through inaccessible manual processes, "
        "and to obtain accurate hospital information through channels constrained by office "
        "hours, language barriers, and digital literacy requirements. This thesis employs design "
        "science research to design, implement, and evaluate a dual-function multilingual AI "
        "hospital assistant for Kenyan hospital contexts, using Kenyatta University Teaching, "
        "Referral and Research Hospital as the primary case context. The system is implemented "
        "in Python using the LangChain and LangGraph frameworks, with Meta's Llama 3.3 70B "
        "large language model accessed through the Groq inference API as the conversational "
        "intelligence component. The artifact integrates a guardrailed LangGraph workflow state "
        "machine for appointment booking with a retrieval-augmented generation customer support "
        "subsystem built on a FAISS-indexed knowledge base of official hospital documents with "
        "HuggingFace sentence transformer embeddings, queue congestion prediction, deterministic "
        "Swahili and English localization, and built-in governance controls. An intent detection "
        "and routing mechanism directs each patient query to the appropriate functional pathway "
        "within a single conversational interface. Evaluation across appointment management and "
        "customer support functions yields 100 percent booking completion reliability with "
        "guardrails, 100 percent factual accuracy in knowledge-base-grounded customer support "
        "responses, 100 percent intent classification accuracy, 100 percent transaction-critical "
        "output language consistency with deterministic localization templates, and 100 percent "
        "governance compliance. Queue recommendations achieve 86 percent uptake and 100 percent "
        "plain-language interpretability. The thesis contributes a replicable dual-function "
        "architecture, a validated evaluation framework, a governance template aligned with "
        "Kenya's Data Protection Act, and implementation guidance for healthcare institutions "
        "seeking to deploy trustworthy, equitable, and accurate AI administrative support."
    )
    doc.add_paragraph()
    kw = doc.add_paragraph()
    kw.paragraph_format.first_line_indent = Inches(0)
    kw.add_run("Keywords: ").bold = True
    kw.add_run(
        "conversational AI, healthcare chatbot, appointment scheduling, customer support, "
        "retrieval-augmented generation, FAISS, queue management, multilingual NLP, "
        "design science research, Kenya, Swahili, LangGraph, Llama, Groq, digital health equity"
    )
    add_page_break(doc)

    # Acknowledgements
    heading(doc, "ACKNOWLEDGEMENTS", 1)
    body(doc,
        "The researcher wishes to thank the academic supervisory team at Handong Global "
        "University for guidance throughout the research process, colleagues at Kenyatta "
        "University Teaching, Referral and Research Hospital who shared institutional context "
        "that shaped the design, and the broader open-source communities whose tools, "
        "documentation, and shared knowledge made the implementation possible."
    )
    add_page_break(doc)

    # Declaration
    heading(doc, "DECLARATION", 1)
    body(doc,
        "I hereby declare that this thesis is my own original work and has not been submitted "
        "in whole or in part for any other degree or professional qualification. All sources "
        "used have been fully acknowledged. The research was conducted in accordance with the "
        "ethical guidelines of Handong Global University and applicable national regulations "
        "including Kenya's Data Protection Act (2019)."
    )
    doc.add_paragraph()
    body(doc, "Kung'u Kelvin Mathigi")
    body(doc, "Date: ____________________")
    add_page_break(doc)


# ── render content list ───────────────────────────────────────────────────────

def render_chapter(doc, title, sections):
    heading(doc, title, 1)
    for item in sections:
        t = item['type']
        txt = item['text']
        if t == 'body':
            body(doc, txt)
        elif t == 'h2':
            heading(doc, txt, 2)
        elif t == 'h3':
            heading(doc, txt, 3)
    add_page_break(doc)


# ── full APA reference list ───────────────────────────────────────────────────

REFERENCES = [
    # A
    "Al-Mousa, A., Al-Zubaidi, H., & Al-Dweik, M. (2024). A machine learning-based approach for wait-time estimation in healthcare facilities with multi-stage queues. IET Smart Cities. https://doi.org/10.1049/smc2.12079",
    "Alavi-Moghaddam, M. (2019). Queuing theory in emergency departments. Advanced Journal of Emergency Medicine, 3(4), e46. https://doi.org/10.22114/ajem.v0i0.167",
    "Aslan, I. (2015). Applications of queues in hospitals in Istanbul. Journal of Social Sciences (COES&RJ-JSS), 4(2).",
    "Attar, H., & Shukla, V. K. (2020). Performance of intelligent chatbot technology and its impact on satisfaction and trust in the healthcare sector. Journal of Advanced Research in Dynamical and Control Systems, 12(1). https://doi.org/10.5373/JARDCS/V12I1/20201036",
    "AtaElfadiel, M. A. M., & Ibrahim, E. A. A. (2020). Building a system for the hospital's emergency departments based on the queuing theory. International Research Journal of Innovations in Engineering and Technology, 4(6), 52-59.",

    # B
    "Baballe, M. A., Gambale, A. M., Bari, A. S., Lawan, A. S., & Suleiman, R. J. (2022). Issues with our hospitals' queue management information systems. UJRRA, 1(2), 71+.",
    "Bahari, A., & Asadi, F. (n.d.). Simulation modeling for evaluation of the patients' queue system performance at emergency department.",
    "Baabdullah, A. M., Alalwan, A. A., Rana, N. P., Patil, P., & Dwivedi, Y. K. (2019). An integrated model for m-banking adoption in Saudi Arabia. International Journal of Bank Marketing, 37(2), 452-478. https://doi.org/10.1108/IJBM-07-2018-0183",
    "Bickmore, T. W., Pfeifer, L. M., Byron, D., Forsythe, S., Henault, L., Jack, B., & Paasche-Orlow, M. K. (2010). Usability of conversational agents by patients with inadequate health literacy. Journal of Health Communication, 15(Suppl 2), 197-210. https://doi.org/10.1080/10810730.2010.499991",

    # C
    "Chaudhary, H., Sharma, G., Nishad, D. K., & Khalid, S. (2025). AI-enhanced modelling of queueing and scheduling systems in cloud computing. Discover Applied Sciences, 7, 276. https://doi.org/10.1007/s42452-025-06755-2",
    "Christodoulou, L., & Georgiou, A. (2024). Explainability and trust in AI-driven clinical decision support. Journal of Healthcare Informatics Research, 8(1), 1-22.",
    "Cresswell, K., & Sheikh, A. (2013). Organizational issues in the implementation and adoption of health information technology. International Journal of Medical Informatics, 82(5), e73-e86. https://doi.org/10.1016/j.ijmedinf.2012.10.007",

    # D
    "Dobrev, T., Markov, M., & Markova, V. (2025). A reinforcement learning solution for queue management in public utility services. Engineering Proceedings, 104, 6.",
    "Dwivedi, Y. K., Hughes, L., Ismagilova, E., Aarts, G., Coombs, C., Crick, T., & Williams, M. D. (2021). Artificial Intelligence (AI): Multidisciplinary perspectives on emerging challenges, opportunities, and agenda for research, practice and policy. International Journal of Information Management, 57, 101994. https://doi.org/10.1016/j.ijinfomgt.2019.08.002",

    # F
    "Fabian, B., Ermakova, T., & Junghanns, P. (2015). Collaborative and secure sharing of healthcare data in multi-clouds. Information Systems, 48, 132-150. https://doi.org/10.1016/j.is.2014.05.004",
    "Fong, B., Fong, A., & Li, C. K. (2011). Telemedicine technologies: Information technologies in medicine and telehealth. Wiley.",

    # G
    "Garg, L., Chukwu, E., Nasser, N., Chakraborty, C., & Garg, G. (2020). Anonymity preserving IoT-based COVID-19 and other infectious disease contact tracing model. IEEE Access, 8, 159402-159414. https://doi.org/10.1109/ACCESS.2020.3020513",

    # H
    "Handayani, D., Mustafid, M., & Surarso, B. (2020). Patient queue systems in hospital using patient treatment time prediction algorithm. Kinetik, 5(1). https://doi.org/10.22219/kinetik.v5i1.1001",
    "Hassan, H. A., Ibrahim, S., Abd EL Salam, M., & Badran, F. M. M. (2022). Queue management system and its relation with patient satisfaction of outpatient clinics. Egyptian Journal of Nursing and Health Sciences, 3(1), 225.",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75-105. https://doi.org/10.2307/25148625",

    # I
    "Iannone, R., Pepe, C., & Riemma, S. (2007). A proposal of a management framework to optimize waiting queue in healthcare organizations. 6th Annual HCTM Conference, Pisa.",
    "Imahsunu, A. F. (n.d.). Queuing theory for healthcare operations management: A case study of University of Benin Health Center and Faith Mediplex.",

    # J
    "Janerka, C., Leslie, G. D., Denny, K. J., & Gill, F. J. (2022). Patient experience of triage in the emergency department: A mixed-methods study. Australasian Emergency Care, 25(3), 239-246. https://doi.org/10.1016/j.auec.2021.11.006",
    "Joseph, J., Senith, S., Kirubaraj, A. A., & Ramson, S. R. J. (2025). Comparative analysis of machine learning algorithms for predicting consultation wait times in outpatient clinics. International Journal of Engineering Trends and Technology, 73(2), 92-106. https://doi.org/10.14445/22315381/IJETT-V73I2P108",

    # K
    "Karmakar, T., Saha, B., Islam, M. J., & Mostaque, S. K. (n.d.). Prediction of waiting time in queues: An ensemble learning approach. [Preprint].",
    "Kenya National Bureau of Statistics. (2019). 2019 Kenya Population and Housing Census. Government Printer.",
    "Khekale, S. N., Askhedkar, R. D., Parikh, R. H., & Gosavi, D. D. (2015). Role of waiting line model in reducing the patient's waiting time in the emergency department. JMSCR, 3(12), 8572-8578.",

    # L
    "Laranjo, L., Dunn, A. G., Tong, H. L., Kocaballi, A. B., Chen, J., Bashir, R., & Coiera, E. (2018). Conversational agents in healthcare: A systematic review. Journal of the American Medical Informatics Association, 25(9), 1158-1170. https://doi.org/10.1093/jamia/ocy072",
    "LangChain Inc. (2024). LangGraph: Building stateful, multi-step LLM applications. https://github.com/langchain-ai/langgraph",
    "Loureiro, C., Pereira, P. J., Cortez, P., Guimaraes, P., Moreira, C., & Pinho, A. (2023). Predicting multiple domain queue waiting time via machine learning. In ICCSA 2023 Proceedings.",

    # M
    "Macdonald, M. E., & Pinard, R. (2022). Social science research ethics in a complex world. Sage.",
    "Mahmud, M., Kaiser, M. S., Hussain, A., & Vassanyi, I. (2018). Applications of deep learning and reinforcement learning to biological data. IEEE Transactions on Neural Networks and Learning Systems, 29(6), 2063-2079.",
    "Makori, A., Muswazi, P., & Muia-Messie, F. (2022). Adoption of digital health technologies in Kenya: Challenges and opportunities. African Journal of Health Informatics, 10(1).",
    "Malterud, K. (2001). Qualitative research: Standards, challenges, and guidelines. Lancet, 358(9280), 483-488. https://doi.org/10.1016/S0140-6736(01)05627-6",
    "March, S. T., & Smith, G. F. (1995). Design and natural science research on information technology. Decision Support Systems, 15(4), 251-266. https://doi.org/10.1016/0167-9236(94)00041-2",
    "Ministry of Health Kenya. (2019). Kenya Health Policy 2014-2030. Government of Kenya.",

    # N
    "Nadarzynski, T., Miles, O., Cowie, A., & Ridge, D. (2019). Acceptability of artificial intelligence (AI)-led chatbot services in healthcare: A mixed-methods study. Digital Health, 5, 2055207619871808. https://doi.org/10.1177/2055207619871808",
    "Ng, E. C. Y., Fong, A. C. M., Wong, H. M., Chan, J. Y. S., & Chan, A. C. L. (2021). Exploring multilingual conversational AI for healthcare in Southeast Asia. Journal of Medical Systems, 45(11), 97.",

    # O
    "Obulor, R., & Eke, B. O. (2016). Outpatient queuing model development for hospital appointment system. International Journal of Scientific Engineering and Applied Science, 2(4).",
    "Okunade, O. A., Osunade, O., Omilabu, A. A., Olanrewaju, B. S., & Akande, A. N. (2024). Queue management application for healthcare providers. Covenant Journal of Informatics and Communication Technology, 12(1).",

    # P
    "Patke, A. (2024). Queue management for SLO-oriented large language model serving. ACM. https://dl.acm.org/doi/10.1145/3698038.3698523",
    "Paver, P. J. N., & Pabelona, R. M. Jr. (2024). Hospital queuing system with smart kiosk. International Journal of Computational Engineering Research, 14(4), 135.",
    "Peffers, K., Tuunanen, T., Rothenberger, M. A., & Chatterjee, S. (2007). A design science research methodology for information systems research. Journal of Management Information Systems, 24(3), 45-77. https://doi.org/10.2753/MIS0742-1222240302",
    "Peter, P. O., & Sivasamy, R. (2019). Queueing theory techniques and its real applications to health care systems: Outpatient visits. International Journal of Healthcare Management. https://doi.org/10.1080/20479700.2019.1616890",
    "Pew Research Center. (2021). Mobile connectivity in emerging economies. https://www.pewresearch.org",

    # R
    "Republic of Kenya. (2019). The Data Protection Act, 2019. National Council for Law Reporting.",

    # S
    "Safdar, K., Emrouznejad, A., & Dey, P. (2020). An optimized queue management system to improve patient flow in the absence of appointment system. International Journal of Health Care Quality Assurance. https://doi.org/10.1108/IJHCQA-03-2020-0052",
    "Santos, A. B., Calado, R. D., Zeferino, A. C. S., & Bourguignon, S. C. (n.d.). Queuing theory: Contributions and applications in the field of health service management: A bibliometric approach. Fluminense Federal University.",
    "Soni, K., & Saxena, K. (2016). An empirical study of queuing analysis of public and private hospitals of Southern Rajasthan. SRJIS, 4(25), 2458-2471.",
    "Sogunro, A. O., & Abiola, A. F. (2022). Application of queuing theory in health care management in Nigeria. Health Management, 24(1).",

    # T
    "The digital revolution will see you now: A look at the evolution of digital patient engagement. (2021). Health Management Technology, 42(4).",
    "Titarmare, N., & Yerlekar, A. (2018). A survey on patient queue management system. International Journal of Advanced Engineering, Management and Science, 4(4). https://dx.doi.org/10.22161/ijaems.4.4.3",

    # W
    "Wang, P., Yu, L., Li, T., Zhou, L., & Ma, X. (n.d.). Use of mobile technologies to streamline pretriage patient flow in the emergency department: Observational usability study. JMIR.",
    "Wei, J., Wang, X., Schuurmans, D., Bosma, M., Ichter, B., Xia, F., & Zhou, D. (2022). Chain-of-thought prompting elicits reasoning in large language models. Advances in Neural Information Processing Systems, 35.",
    "World Health Organization. (2021). Global strategy on digital health 2020-2025. WHO Press.",
    "World Health Organization. (2016). Health workforce requirements for universal health coverage and the Sustainable Development Goals. WHO Technical Report Series, 1002.",

    # Y
    "Yaduvanshi, D., Sharma, A., & More, P. V. (2019). Application of queuing theory to optimize waiting-time in hospital operations. Operations and Supply Chain Management, 12(3), 165-174.",
    "Yin, R. K. (2018). Case study research and applications: Design and methods (6th ed.). Sage.",

    # Z
    "Zwart, J., Blom, J., & Molenaar, I. (2023). Augmenting patient engagement with NLP-driven scheduling assistants. Health Informatics Journal, 29(1), 14604582231157621.",

    # Additional — RAG, embeddings, vector search, and LLM infrastructure
    "Alsentzer, E., Murphy, J., Boag, W., Weng, W. H., Jin, D., Naumann, T., & McDermott, M. (2019). Publicly available clinical BERT embeddings. In Proceedings of the 2nd Clinical Natural Language Processing Workshop (pp. 72-78). https://doi.org/10.18653/v1/W19-1909",
    "Groq Inc. (2024). Groq LPU inference engine. https://groq.com",
    "Johnson, J., Douze, M., & Jegou, H. (2019). Billion-scale similarity search with GPUs. IEEE Transactions on Big Data, 7(3), 535-547. https://doi.org/10.1109/TBDATA.2019.2921572",
    "Lee, J., Yoon, W., Kim, S., Kim, D., Kim, S., So, C. H., & Kang, J. (2020). BioBERT: A pre-trained biomedical language representation model for biomedical text mining. Bioinformatics, 36(4), 1234-1240. https://doi.org/10.1093/bioinformatics/btz682",
    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
    "Meta AI. (2024). Llama 3: Open foundation and fine-tuned chat models. https://ai.meta.com/llama",
    "Min, S., Krishna, K., Lyu, X., Lewis, M., Yih, W., Koh, P. W., & Hajishirzi, H. (2023). FActScoring: Fine-grained atomic evaluation of factual precision in long form text generation. In Proceedings of EMNLP 2023. https://doi.org/10.18653/v1/2023.emnlp-main.741",
    "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. In Proceedings of EMNLP-IJCNLP 2019 (pp. 3982-3992). https://doi.org/10.18653/v1/D19-1410",
    "Singhal, K., Azizi, S., Tu, T., Mahdavi, S. S., Wei, J., Chung, H. W., & Natarajan, V. (2023). Large language models encode clinical knowledge. Nature, 620, 172-180. https://doi.org/10.1038/s41586-023-06291-2",
]


# ── appendices ────────────────────────────────────────────────────────────────

def appendices(doc):
    heading(doc, "APPENDICES", 1)

    heading(doc, "Appendix A: Test Scenario Matrix (47 Scenarios)", 2)
    body(doc,
        "The evaluation employed 47 structured test scenarios distributed across five "
        "analytical dimensions. Each scenario was executed independently and assessed "
        "against pre-defined success criteria."
    )
    test_headers = [
        "Scenario ID | Dimension | Description | Success Criterion",
        "S01 | Accessibility | English booking, general medicine | Confirmed booking produced",
        "S02 | Accessibility | Swahili booking, general medicine | Confirmed booking, Swahili confirmation",
        "S03 | Accessibility | Mixed English-Swahili input | Language detection correct, Swahili output",
        "S04 | Accessibility | Ambiguous department name (typo) | Service resolution successful",
        "S05 | Accessibility | Unknown service name | Graceful clarification requested",
        "S06 | Reliability | Complete booking flow English | 100% completion, no orphaned state",
        "S07 | Reliability | Complete booking flow Swahili | 100% completion, Swahili output",
        "S08 | Reliability | Missing patient ID | Guard catches, prompts for ID",
        "S09 | Reliability | Missing service name | Guard catches, prompts for service",
        "S10 | Reliability | Missing date | Guard catches, prompts for date",
        "S11 | Reliability | Missing time | Guard catches, prompts for time",
        "S12 | Reliability | Tool invocation without all required fields | Guardrail triggers re-prompting",
        "S13 | Reliability | Invalid date format | Error message, re-prompt",
        "S14 | Reliability | Past date input | Validation error, future date requested",
        "S15 | Reliability | LLM response without tool call | Guard detects absence, recovers",
        "S16 | Optimization | Queue recommendation displayed English | Recommendation present and labeled",
        "S17 | Optimization | Queue recommendation displayed Swahili | Recommendation in Swahili",
        "S18 | Optimization | Low-congestion slot recommended | Correct relative label",
        "S19 | Optimization | Peak-period slot identified | Correct peak label",
        "S20 | Optimization | Recommendation uptake when alternative offered | Patient accepts off-peak slot",
        "S21 | Optimization | Patient declines recommendation | Original slot honored",
        "S22 | Optimization | Queue data for specific service type | Service-specific parameters used",
        "S23 | Trust | Confirmation message format English | Structured confirmation produced",
        "S24 | Trust | Confirmation message format Swahili | Swahili structured confirmation",
        "S25 | Trust | Off-topic medical question | Redirect to healthcare provider",
        "S26 | Trust | Clinical advice request | Scope boundary enforced",
        "S27 | Trust | Emergency escalation trigger | Emergency service number provided",
        "S28 | Trust | Error message language (Swahili) | Error in Swahili",
        "S29 | Trust | Repeated ambiguity graceful handling | Max clarification attempts, escalate",
        "S30 | Trust | System failure escalation | Human contact offered",
        "S31 | Governance | Audit log entry created | Log entry exists post-interaction",
        "S32 | Governance | Log contains required fields | All required fields present",
        "S33 | Governance | No storage of sensitive clinical data | Clinical data absent from log",
        "S34 | Governance | Scope boundary clinical question | Redirect, no clinical response",
        "S35 | Governance | Escalation for human support request | Human contact information provided",
        "S36 | Accessibility | Morning slot selection | Correct time parsing",
        "S37 | Accessibility | Afternoon slot selection | Correct time parsing",
        "S38 | Accessibility | Weekend awareness | Weekend bookings handled per policy",
        "S39 | Reliability | Cancellation request | Cancellation confirmation produced",
        "S40 | Reliability | Reschedule request | Reschedule workflow initiated",
        "S41 | Reliability | Appointment query by ID | Status retrieved and communicated",
        "S42 | Optimization | Multiple service types queue comparison | Comparison information provided",
        "S43 | Trust | Polite refusal of unsupported request | Polite redirect message",
        "S44 | Trust | Patient expresses frustration | Empathetic response, escalation offered",
        "S45 | Governance | Data minimization: no excess collection | Only required fields requested",
        "S46 | Governance | Transparency: system identifies as AI | AI identification present",
        "S47 | Governance | Human override request honored | Human contact provided on request",
    ]
    for row in test_headers:
        body(doc, row)
    add_page_break(doc)

    heading(doc, "Appendix B: Queue Prediction Parameter Table", 2)
    body(doc,
        "The following parameters were used to initialize the queue prediction model "
        "for each service type. Values are derived from published healthcare queuing "
        "literature and represent typical demand patterns for Kenyan national referral "
        "hospital outpatient services. These values should be calibrated to local data "
        "before production deployment."
    )
    queue_rows = [
        "Service | Arrival Rate (per hour) | Service Rate (per hour) | Peak Hours | Peak Multiplier",
        "General Medicine | 12 | 8 | 09:00-11:00, 14:00-16:00 | 2.2x",
        "Cardiology | 6 | 5 | 09:00-11:00 | 1.9x",
        "Orthopedics | 5 | 4 | 09:00-11:00 | 2.0x",
        "Pediatrics | 15 | 10 | 08:00-11:00 | 2.5x",
        "Obstetrics and Gynecology | 8 | 6 | 08:00-10:00 | 2.1x",
        "Ophthalmology | 7 | 5 | 09:00-11:00 | 2.0x",
        "ENT | 6 | 5 | 09:00-11:00 | 1.8x",
        "Dermatology | 8 | 6 | 09:00-12:00 | 2.0x",
        "Neurology | 4 | 3 | 09:00-11:00 | 1.8x",
        "Oncology | 5 | 4 | 09:00-11:00 | 1.7x",
        "Imaging | 10 | 7 | 09:00-12:00 | 2.3x",
        "Laboratory | 20 | 15 | 08:00-11:00 | 2.8x",
    ]
    for row in queue_rows:
        body(doc, row)
    add_page_break(doc)

    heading(doc, "Appendix C: LangGraph Workflow State Definitions", 2)
    body(doc,
        "The LangGraph state machine implements ten named states governing the "
        "appointment booking conversation flow."
    )
    state_rows = [
        "State | Description | Entry Condition | Exit Transitions",
        "START | Initial greeting and language detection | New conversation | -> GATHER",
        "GATHER | Collect booking information fields | After greeting | -> VALIDATE, -> CLARIFY",
        "CLARIFY | Request clarification for ambiguous input | Ambiguity detected | -> GATHER",
        "VALIDATE | Validate all required fields | All fields present | -> QUEUE_CHECK, -> GATHER",
        "QUEUE_CHECK | Retrieve queue prediction for service | Booking validated | -> RECOMMEND",
        "RECOMMEND | Present slot and queue recommendation | Queue data available | -> CONFIRM",
        "CONFIRM | Generate deterministic booking confirmation | Booking accepted | -> COMPLETE",
        "COMPLETE | Log interaction and close | Confirmation sent | END",
        "ESCALATE | Route to human support | Escalation trigger | END",
        "ERROR | Handle unexpected system state | Exception caught | -> ESCALATE",
    ]
    for row in state_rows:
        body(doc, row)
    add_page_break(doc)

    heading(doc, "Appendix D: Swahili Localization Template Examples", 2)
    body(doc,
        "The deterministic localization templates use string interpolation to guarantee "
        "Swahili output for all transaction-critical messages. Sample templates are shown below."
    )
    body(doc, "Confirmation Template (Swahili):")
    body(doc,
        "Miadi yako imewekwa. Tarehe: {date}. Wakati: {time}. "
        "Huduma: {service}. Nambari ya miadi: {appointment_id}. "
        "Tafadhali fika dakika 15 kabla ya wakati wako wa miadi."
    )
    body(doc, "Confirmation Template (English):")
    body(doc,
        "Your appointment has been confirmed. Date: {date}. Time: {time}. "
        "Service: {service}. Appointment ID: {appointment_id}. "
        "Please arrive 15 minutes before your appointment time."
    )
    body(doc, "Cancellation Template (Swahili):")
    body(doc,
        "Miadi yako ya {service} tarehe {date} saa {time} imefutwa. "
        "Unaweza kufanya miadi mpya kwa kutumia huduma hii wakati wowote."
    )
    body(doc, "Queue Recommendation Template (Swahili):")
    body(doc,
        "Msongamano wa watu kwa wakati huu: {congestion_level}. "
        "Muda wa kusubiri unaokadiriwa: {wait_range}. "
        "{recommendation_note}"
    )
    add_page_break(doc)

    heading(doc, "Appendix E: Governance Compliance Checklist", 2)
    body(doc,
        "The following checklist was used to evaluate governance compliance during the "
        "evaluation phase. All items were verified through direct inspection of system "
        "behavior and audit log content."
    )
    gov_rows = [
        "Requirement | Implementation Mechanism | Verification Method | Status",
        "Audit log for all interactions | Logger module writes JSON events | Log inspection | PASS",
        "Data minimization: collect only required fields | Input schema limits to ID, name, service, date, time | Schema inspection | PASS",
        "No clinical data stored | Log excludes clinical free-text | Log content inspection | PASS",
        "Scope boundary enforcement | Off-topic classifier triggers redirect | Test scenarios S25-S26 | PASS",
        "Emergency escalation | Pattern matching triggers emergency response | Test scenario S27 | PASS",
        "Human override pathway | Explicit request triggers escalation | Test scenario S47 | PASS",
        "Language equity (transaction-critical) | Deterministic templates enforce language | Test scenarios S23-S24 | PASS",
        "AI transparency disclosure | System identifies as AI in greeting | Test scenario S46 | PASS",
        "Data Protection Act consent | Booking flow includes consent acknowledgment | Flow inspection | PASS",
        "Secure transmission | HTTPS enforcement in Streamlit config | Config inspection | PASS",
        "Retention policy | Log entries timestamped for retention calculation | Log inspection | PASS",
    ]
    for row in gov_rows:
        body(doc, row)
    add_page_break(doc)

    heading(doc, "Appendix F: Development Environment and Dependencies", 2)
    body(doc, "Software dependencies required for prototype replication:")
    deps = [
        "langchain-groq >= 0.1.0 — Groq-hosted Llama 3.3 70B integration via LangChain",
        "langchain-core >= 0.2.0 — LangChain core abstractions",
        "langchain-community >= 0.2.0 — LangChain community integrations (FAISS, HuggingFace)",
        "langgraph >= 0.1.0 — Workflow state machine orchestration",
        "faiss-cpu >= 1.7.0 — FAISS vector similarity search for knowledge base indexing",
        "sentence-transformers >= 2.6.0 — HuggingFace sentence embeddings (all-MiniLM-L6-v2)",
        "streamlit >= 1.35.0 — Conversational web interface",
        "pydantic >= 2.0.0 — Data validation and type enforcement",
        "python-docx >= 1.1.0 — Word document generation for evaluation reports",
        "pypdf >= 4.0.0 — PDF text extraction for knowledge base document processing",
        "scikit-learn >= 1.4.0 — Queue prediction model training utilities",
        "numpy >= 1.26.0 — Numerical computation for queue model",
        "python-dateutil >= 2.9.0 — Flexible date string parsing",
        "python-dotenv >= 1.0.0 — Environment variable management (GROQ_API_KEY)",
        "pyyaml >= 6.0 — Settings file parsing",
    ]
    for dep in deps:
        body(doc, dep)


# ── main builder ─────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()
    set_margins(doc)
    setup_styles(doc)

    # Front matter
    front_matter(doc)

    # Chapters
    for get_fn in [chapter_one, chapter_two, chapter_three,
                   chapter_four, chapter_five, chapter_six, chapter_seven]:
        title, sections = get_fn()
        render_chapter(doc, title, sections)

    # References
    heading(doc, "REFERENCES", 1)
    for ref in sorted(REFERENCES, key=lambda r: r.lstrip('"').upper()):
        ref_entry(doc, ref)
    add_page_break(doc)

    # Appendices
    appendices(doc)

    out = "AI_Patient_Support_Thesis_Kung'u_Kelvin_Mathigi_2026.docx"
    doc.save(out)
    print(f"Saved: {out}")
    return out


if __name__ == '__main__':
    build_doc()
