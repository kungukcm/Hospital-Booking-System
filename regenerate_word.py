from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read master manuscript
manuscript_path = Path('thesis_manuscript/MASTER_THESIS_MANUSCRIPT.md')
with open(manuscript_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Split content into lines
lines = content.split('\n')

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    
    # Skip separator lines
    if line.strip() == '---':
        i += 1
        continue
    
    # Skip empty lines
    if not line.strip():
        i += 1
        continue
    
    # Process headings
    if line.startswith('#'):
        heading_level = len(line) - len(line.lstrip('#'))
        heading_text = line.lstrip('#').strip()
        
        if heading_level == 1:
            p = doc.add_heading(heading_text, level=1)
        elif heading_level == 2:
            p = doc.add_heading(heading_text, level=2)
        elif heading_level == 3:
            p = doc.add_heading(heading_text, level=3)
        else:
            p = doc.add_heading(heading_text, level=4)
    
    # Process paragraphs
    else:
        p = doc.add_paragraph(line)
    
    i += 1

# Save document
output_path = Path('thesis_manuscript/FINAL_THESIS_DOCUMENT.docx')
doc.save(str(output_path))

print(f'✓ FINAL_THESIS_DOCUMENT.docx regenerated successfully')
print(f'✓ Document saved to thesis_manuscript/FINAL_THESIS_DOCUMENT.docx')
